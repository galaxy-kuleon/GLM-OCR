#!/usr/bin/env python3
"""dsl_to_docx.py - Convert per-page XML DSL files into a single DOCX document.

Usage:
    python dsl_to_docx.py --workspace PATH --output PATH [--dsl-dir DIR_NAME]

Reads all $WORKSPACE/<dsl-dir>/page-*.xml files, sorted by page number,
and produces a single DOCX file.
"""

import argparse
import glob
import re
import sys
from pathlib import Path

from lxml import etree
from docx import Document
from docx.shared import Pt, Emu, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls


# ---------------------------------------------------------------------------
# Font mapping for font-family detection
# ---------------------------------------------------------------------------

FONT_MAPPING = {
    "serif": {"latin": "Times New Roman", "cjk": "SimSun"},
    "sans": {"latin": "Arial", "cjk": "SimHei"},
    "mono": {"latin": "Courier New", "cjk": "SimSun"},
}


def get_fonts_for_element(elem, page_font_latin, page_font_cjk):
    """Get appropriate fonts based on element's font-family attribute."""
    ff = elem.get("font-family") if elem is not None else None

    if ff:
        mapping = FONT_MAPPING.get(ff, FONT_MAPPING["serif"])
        return mapping["latin"], mapping["cjk"]

    return page_font_latin, page_font_cjk


# ---------------------------------------------------------------------------
# Math/Equation OMML conversion
# ---------------------------------------------------------------------------


def latex_to_omml(latex: str) -> str:
    """
    Convert simple LaTeX to OMML XML string.
    This is a simplified converter for common patterns.
    """
    if not latex:
        return ""

    # Remove inline math delimiters
    latex = latex.strip()
    if latex.startswith("$") and latex.endswith("$"):
        latex = latex[1:-1].strip()

    # Handle fractions: \frac{a}{b}
    frac_pattern = r"\\frac\s*\{\s*([^}]+)\s*\}\s*\{\s*([^}]+)\s*\}"

    def replace_frac(match):
        num = match.group(1)
        den = match.group(2)
        return f"""<m:f>
            <m:num><m:r><m:t>{num}</m:t></m:r></m:num>
            <m:den><m:r><m:t>{den}</m:t></m:r></m:den>
        </m:f>"""

    omml = re.sub(frac_pattern, replace_frac, latex)

    # Handle superscripts: x^{n} or x^n
    sup_pattern = r"([a-zA-Z0-9])\^\{\s*([^}]+)\s*\}"

    def replace_sup(match):
        base = match.group(1)
        exp = match.group(2)
        return f"""<m:sSup>
            <m:e><m:r><m:t>{base}</m:t></m:r></m:e>
            <m:sup><m:r><m:t>{exp}</m:t></m:r></m:sup>
        </m:sSup>"""

    omml = re.sub(sup_pattern, replace_sup, omml)

    # Handle subscripts: x_{n} or x_n
    sub_pattern = r"([a-zA-Z0-9])_\{\s*([^}]+)\s*\}"

    def replace_sub(match):
        base = match.group(1)
        sub = match.group(2)
        return f"""<m:sSub>
            <m:e><m:r><m:t>{base}</m:t></m:r></m:e>
            <m:sub><m:r><m:t>{sub}</m:t></m:r></m:sub>
        </m:sSub>"""

    omml = re.sub(sub_pattern, replace_sub, omml)

    # Handle Greek letters
    greek_map = {
        r"\\alpha": "α",
        r"\\beta": "β",
        r"\\gamma": "γ",
        r"\\delta": "δ",
        r"\\pi": "π",
        r"\\sigma": "σ",
        r"\\mu": "μ",
        r"\\lambda": "λ",
        r"\\theta": "θ",
        r"\\phi": "φ",
        r"\\omega": "ω",
    }

    for latex_greek, unicode_greek in greek_map.items():
        omml = omml.replace(latex_greek, unicode_greek)

    # Handle operators
    op_map = {
        r"\\times": "×",
        r"\\div": "÷",
        r"\\pm": "±",
        r"\\leq": "≤",
        r"\\geq": "≥",
        r"\\infty": "∞",
        r"\\cdot": "·",
        r"\\partial": "∂",
        r"\\nabla": "∇",
        r"\\sum": "∑",
        r"\\int": "∫",
        r"\\prod": "∏",
    }

    for latex_op, unicode_op in op_map.items():
        omml = omml.replace(latex_op, unicode_op)

    # If no special handling applied, wrap as plain text
    if omml == latex:
        omml = f"<m:r><m:t>{latex}</m:t></m:r>"

    return omml


def add_math_equation(paragraph, latex_expr: str):
    """
    Convert LaTeX expression to Word OMML (Office Math Markup Language)
    and add to paragraph.
    """
    # Create math element
    math_elem = OxmlElement("m:oMath")

    # Parse and convert LaTeX to OMML
    omml = latex_to_omml(latex_expr)

    # Parse OMML XML string and append
    if omml:
        try:
            omml_elem = parse_xml(
                f'<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math">{omml}</m:oMath>'
            )
            paragraph._p.append(omml_elem)
        except Exception:
            # Fallback: add as plain text
            run = paragraph.add_run(latex_expr)
            run.font.name = "Cambria Math"

    return paragraph


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def pts_to_emu(pts):
    return int(float(pts) * 914400 / 72)


def parse_color_rgb(s):
    """Parse 'R,G,B' string into (r, g, b) tuple."""
    if not s:
        return None
    parts = [int(x.strip()) for x in s.split(",")]
    if len(parts) == 3:
        return tuple(min(255, max(0, v)) for v in parts)
    return None


def parse_bool(val, default=False):
    if val is None:
        return default
    return val.lower() in ("true", "1", "yes")


def set_east_asian_font(run, font_name):
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f"<w:rFonts {nsdecls('w')}/>")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def has_cjk(text):
    return any("\u4e00" <= c <= "\u9fff" or "\u3000" <= c <= "\u303f" for c in text)


def apply_run_attrs(run, elem, page_font_latin, page_font_cjk):
    """Apply XML <run> attributes to a python-docx Run object."""
    fs = elem.get("font-size-pt")
    if fs:
        run.font.size = Pt(float(fs))

    if parse_bool(elem.get("bold")):
        run.bold = True
    if parse_bool(elem.get("italic")):
        run.italic = True
    if parse_bool(elem.get("underline")):
        run.font.underline = True
    if parse_bool(elem.get("superscript")):
        run.font.superscript = True

    color = parse_color_rgb(elem.get("color-rgb"))
    if color:
        run.font.color.rgb = RGBColor(*color)

    # Get font family from element or use page defaults
    latin_font, cjk_font = get_fonts_for_element(elem, page_font_latin, page_font_cjk)

    font_name = elem.get("font-name")
    text = elem.text or ""
    if font_name:
        run.font.name = font_name
        if has_cjk(text):
            set_east_asian_font(run, font_name)
    else:
        run.font.name = latin_font
        if has_cjk(text):
            set_east_asian_font(run, cjk_font)


def get_alignment(val):
    if not val:
        return None
    m = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return m.get(val.lower())


def add_page_break(doc):
    para = doc.add_paragraph()
    run_obj = para.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run_obj._r.append(br)
    return para


def set_cell_shading(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    tcPr.append(shading)


def set_paragraph_shading(paragraph, color_hex):
    pPr = paragraph._p.get_or_add_pPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    pPr.append(shading)


def set_run_shading(run, color_hex):
    rPr = run._r.get_or_add_rPr()
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>'
    )
    rPr.append(shading)


def set_cell_margins(cell, top=0, bottom=0, left=29, right=29):
    """Set cell margins in twips. Default: 0 top/bottom, ~0.5mm left/right."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f"<w:tcMar {nsdecls('w')}>"
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f"</w:tcMar>"
    )
    tcPr.append(margins)


def set_cell_vertical_alignment(cell, align="center"):
    """Set cell vertical alignment. align: 'top', 'center', 'bottom'."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = OxmlElement("w:vAlign")
    vAlign.set(qn("w:val"), align)
    tcPr.append(vAlign)


def set_row_height(row, height_twips, rule="atLeast"):
    """Set table row height. rule: 'atLeast' or 'exact'."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_twips)))
    trHeight.set(qn("w:hRule"), rule)
    trPr.append(trHeight)


def set_cell_borders(cell, color="000000", size=4, style="single"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f"<w:tcBorders {nsdecls('w')}>"
        f'  <w:top w:val="{style}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="{style}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{style}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="{style}" w:sz="{size}" w:space="0" w:color="{color}"/>'
        f"</w:tcBorders>"
    )
    tcPr.append(borders)


def set_invisible_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = parse_xml(
        f"<w:tcBorders {nsdecls('w')}>"
        f'  <w:top w:val="none" w:sz="0" w:space="0"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0"/>'
        f"</w:tcBorders>"
    )
    tcPr.append(borders)


def add_frame_pr(paragraph, x_twips, y_twips, w_twips, h_twips):
    pPr = paragraph._p.get_or_add_pPr()
    frame_pr = OxmlElement("w:framePr")
    frame_pr.set(qn("w:w"), str(int(w_twips)))
    frame_pr.set(qn("w:h"), str(int(h_twips)))
    frame_pr.set(qn("w:hRule"), "exact")
    frame_pr.set(qn("w:hAnchor"), "page")
    frame_pr.set(qn("w:vAnchor"), "page")
    frame_pr.set(qn("w:x"), str(int(x_twips)))
    frame_pr.set(qn("w:y"), str(int(y_twips)))
    frame_pr.set(qn("w:wrap"), "notBeside")
    pPr.insert(0, frame_pr)


def add_paragraph_borders(paragraph, color_hex="000000", size=4):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f"<w:pBdr {nsdecls('w')}>"
        f'<w:top w:val="single" w:sz="{size}" w:space="1" w:color="{color_hex}"/>'
        f'<w:left w:val="single" w:sz="{size}" w:space="1" w:color="{color_hex}"/>'
        f'<w:bottom w:val="single" w:sz="{size}" w:space="1" w:color="{color_hex}"/>'
        f'<w:right w:val="single" w:sz="{size}" w:space="1" w:color="{color_hex}"/>'
        f"</w:pBdr>"
    )
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Element processors
# ---------------------------------------------------------------------------


def process_runs(para, elem, page_font_latin, page_font_cjk):
    """Add <run> children of elem to the paragraph."""
    runs = elem.findall("run")
    if not runs:
        # If no <run> children, use elem text directly
        text = (elem.text or "").strip()
        if text:
            run = para.add_run(text)
            run.font.name = page_font_latin
            run.font.size = Pt(11)
            if has_cjk(text):
                set_east_asian_font(run, page_font_cjk)
        return

    for run_elem in runs:
        text = run_elem.text or ""
        if not text:
            continue
        run = para.add_run(text)
        apply_run_attrs(run, run_elem, page_font_latin, page_font_cjk)


def process_heading(doc, elem, page_font_latin, page_font_cjk):
    level = int(elem.get("level", "1"))
    alignment = get_alignment(elem.get("alignment"))
    bg_color = elem.get("bg-color")

    # Collect text from runs for heading
    runs = elem.findall("run")
    if runs:
        heading_text = "".join((r.text or "") for r in runs)
    else:
        heading_text = (elem.text or "").strip()

    heading = doc.add_heading(heading_text, level=level)

    # Get font family from element
    latin_font, cjk_font = get_fonts_for_element(elem, page_font_latin, page_font_cjk)

    # Get spacing from element
    space_before = elem.get("space-before-pt", "12")
    space_after = elem.get("space-after-pt", "6")
    heading.paragraph_format.space_before = Pt(float(space_before))
    heading.paragraph_format.space_after = Pt(float(space_after))

    if alignment is not None:
        heading.alignment = alignment

    # Apply paragraph background shading
    if bg_color:
        set_paragraph_shading(heading, bg_color)

    # Override heading color to black and apply run styles
    for run in heading.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = latin_font
        if has_cjk(run.text):
            set_east_asian_font(run, cjk_font)

    # Apply specific run styling if present
    if runs:
        for i, run_elem in enumerate(runs):
            if i < len(heading.runs):
                doc_run = heading.runs[i]
                fs = run_elem.get("font-size-pt")
                if fs:
                    doc_run.font.size = Pt(float(fs))
                if parse_bool(run_elem.get("bold")):
                    doc_run.bold = True
                if parse_bool(run_elem.get("italic")):
                    doc_run.italic = True
                color = parse_color_rgb(run_elem.get("color-rgb"))
                if color:
                    doc_run.font.color.rgb = RGBColor(*color)

    return heading


def process_paragraph(doc, elem, page_font_latin, page_font_cjk):
    style_name = elem.get("style")
    bg_color = elem.get("bg-color")

    # Map style attribute to python-docx paragraph styles
    docx_style = None
    if style_name == "list-bullet":
        docx_style = "List Bullet"
    elif style_name == "list-number":
        docx_style = "List Number"

    para = doc.add_paragraph(style=docx_style)

    alignment = get_alignment(elem.get("alignment"))
    if alignment is not None:
        para.alignment = alignment

    space_before = elem.get("space-before-pt", "0")
    space_after = elem.get("space-after-pt", "0")
    line_spacing = elem.get("line-spacing", "1.0")
    para.paragraph_format.space_before = Pt(float(space_before))
    para.paragraph_format.space_after = Pt(float(space_after))
    para.paragraph_format.line_spacing = float(line_spacing)

    # Apply paragraph background shading
    if bg_color:
        set_paragraph_shading(para, bg_color)

    # Check if this is a formula/math paragraph
    is_formula = style_name == "formula"
    if is_formula:
        # Handle math equations
        runs = elem.findall("run")
        for run_elem in runs:
            latex = run_elem.get("latex")
            is_math = run_elem.get("is-math") == "true"

            if is_math and latex:
                # Add as Word equation
                add_math_equation(para, latex)
            else:
                # Regular text run
                text = run_elem.text or ""
                if text:
                    run = para.add_run(text)
                    apply_run_attrs(run, run_elem, page_font_latin, page_font_cjk)
    else:
        process_runs(para, elem, page_font_latin, page_font_cjk)

    return para


def process_table(doc, elem, page_font_latin, page_font_cjk, workspace):
    num_rows = int(elem.get("rows", "0"))
    num_cols = int(elem.get("cols", "0"))
    border_style = elem.get("border-style", "full")

    if num_rows == 0 or num_cols == 0:
        return None

    # Parse col-widths
    col_widths_elem = elem.find("col-widths")
    col_ratios = None
    if col_widths_elem is not None and col_widths_elem.text:
        col_ratios = [float(x.strip()) for x in col_widths_elem.text.split(",")]

    # Calculate table width from bbox
    bbox_str = elem.get("bbox")
    page_width_pts = float(elem.get("page-width-pts", "595"))
    table_width_pts = None
    if bbox_str:
        parts = [float(x) for x in bbox_str.split(",")]
        if len(parts) == 4:
            table_width_pts = (parts[2] - parts[0]) * page_width_pts / 1000

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = "Table Grid"
    table.autofit = False

    # Set column widths
    if table_width_pts and col_ratios and len(col_ratios) == num_cols:
        for i, col in enumerate(table.columns):
            col.width = Pt(table_width_pts * col_ratios[i])
    elif table_width_pts:
        col_w = Pt(table_width_pts / num_cols)
        for col in table.columns:
            col.width = col_w

    # Build occupancy grid
    occupied = [[False] * num_cols for _ in range(num_rows)]

    # Process rows
    for row_elem in elem.findall("row"):
        row_idx = int(row_elem.get("index", "0"))
        is_header = parse_bool(row_elem.get("is-header"))

        col_cursor = 0
        for cell_elem in row_elem.findall("cell"):
            cell_row = int(cell_elem.get("row", str(row_idx)))
            cell_col_attr = cell_elem.get("col")

            if cell_col_attr is not None:
                col_cursor = int(cell_col_attr)
            else:
                while col_cursor < num_cols and occupied[row_idx][col_cursor]:
                    col_cursor += 1

            if col_cursor >= num_cols or row_idx >= num_rows:
                continue

            colspan = int(cell_elem.get("colspan", "1"))
            rowspan = int(cell_elem.get("rowspan", "1"))
            run_children = cell_elem.findall("run")
            if run_children:
                cell_text = "".join((r.text or "") for r in run_children)
            else:
                cell_text = (cell_elem.text or "").strip()
            cell_bold = parse_bool(cell_elem.get("bold"))
            cell_italic = parse_bool(cell_elem.get("italic"))
            cell_alignment = get_alignment(cell_elem.get("alignment"))
            cell_font_size = float(cell_elem.get("font-size-pt", "9"))
            cell_color = parse_color_rgb(cell_elem.get("color-rgb"))
            bg_color = cell_elem.get("bg-color")
            text_bg_color = cell_elem.get("text-bg-color")

            # Mark occupancy
            for mr in range(cell_row, min(cell_row + rowspan, num_rows)):
                for mc in range(col_cursor, min(col_cursor + colspan, num_cols)):
                    occupied[mr][mc] = True

            cell = table.cell(cell_row, col_cursor)

            # Merge if needed
            end_r = min(cell_row + rowspan - 1, num_rows - 1)
            end_c = min(col_cursor + colspan - 1, num_cols - 1)
            if end_r > cell_row or end_c > col_cursor:
                cell.merge(table.cell(end_r, end_c))

            # Tight cell margins for compact layout
            set_cell_margins(cell, top=0, bottom=0, left=29, right=29)
            set_cell_vertical_alignment(cell, "center")

            # Set cell text with line breaks
            cell.paragraphs[0].clear()
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            para.paragraph_format.line_spacing = 1.0
            if cell_alignment is not None:
                para.alignment = cell_alignment

            if run_children:
                # Per-run mixed formatting
                for run_elem in run_children:
                    run_text = run_elem.text or ""
                    if not run_text:
                        continue
                    docx_run = para.add_run(run_text)
                    r_fs = run_elem.get("font-size-pt")
                    docx_run.font.size = Pt(float(r_fs)) if r_fs else Pt(cell_font_size)
                    docx_run.font.name = page_font_latin
                    if has_cjk(run_text):
                        set_east_asian_font(docx_run, page_font_cjk)
                    r_bold = run_elem.get("bold")
                    if r_bold is not None:
                        docx_run.bold = parse_bool(r_bold)
                    elif cell_bold or is_header:
                        docx_run.bold = True
                    r_italic = run_elem.get("italic")
                    if r_italic is not None:
                        docx_run.italic = parse_bool(r_italic)
                    elif cell_italic:
                        docx_run.italic = True
                    r_color = parse_color_rgb(run_elem.get("color-rgb"))
                    if r_color:
                        docx_run.font.color.rgb = RGBColor(*r_color)
                    elif cell_color:
                        docx_run.font.color.rgb = RGBColor(*cell_color)
                    r_text_bg = run_elem.get("text-bg-color")
                    if r_text_bg:
                        set_run_shading(docx_run, r_text_bg)
                    elif text_bg_color:
                        set_run_shading(docx_run, text_bg_color)
            else:
                # Original uniform cell text logic
                lines = cell_text.split("\n") if cell_text else [""]
                for li, line in enumerate(lines):
                    if li > 0:
                        br_run = para.add_run()
                        br_run.add_break()
                    run = para.add_run(line.strip())
                    run.font.size = Pt(cell_font_size)
                    run.font.name = page_font_latin
                    if has_cjk(line):
                        set_east_asian_font(run, page_font_cjk)
                    if cell_bold or is_header:
                        run.bold = True
                    if cell_italic:
                        run.italic = True
                    if cell_color:
                        run.font.color.rgb = RGBColor(*cell_color)
                    if text_bg_color:
                        set_run_shading(run, text_bg_color)

            # Background color
            if bg_color:
                set_cell_shading(cell, bg_color)

            # Borders
            if border_style in ("single", "double", "full"):
                effective_style = "single" if border_style == "full" else border_style
                set_cell_borders(cell, style=effective_style)
            elif border_style == "none":
                set_invisible_cell_borders(cell)

            col_cursor += colspan

    # Set compact row heights based on font size
    for row in table.rows:
        set_row_height(row, height_twips=0, rule="atLeast")

    return table


def process_image(doc, elem, workspace, page_font_latin, page_font_cjk):
    src = elem.get("src", "")
    alignment = get_alignment(elem.get("alignment"))
    bbox_str = elem.get("bbox")
    page_width_pts = float(elem.get("page-width-pts", "595"))

    # Resolve image path
    img_path = Path(workspace) / src
    if not img_path.exists():
        # Try relative to ocr-output
        img_path = Path(workspace) / "ocr-output" / "input" / src
    if not img_path.exists():
        para = doc.add_paragraph(f"[Image missing: {src}]")
        return para

    # Calculate width from bbox
    img_width_inches = None
    if bbox_str:
        parts = [float(x) for x in bbox_str.split(",")]
        if len(parts) == 4:
            width_ratio = (parts[2] - parts[0]) / 1000.0
            usable_width_inches = (page_width_pts / 72.0) - 1.0
            img_width_inches = min(
                width_ratio * usable_width_inches, usable_width_inches
            )

    if img_width_inches and img_width_inches > 0:
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(str(img_path), width=Inches(img_width_inches))
    else:
        doc.add_picture(str(img_path))
        para = doc.paragraphs[-1]

    if alignment is not None:
        para.alignment = alignment

    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after = Pt(1)
    return para


def process_text_frame(doc, elem, page_font_latin, page_font_cjk):
    x_twips = int(elem.get("x-twips", "0"))
    y_twips = int(elem.get("y-twips", "0"))
    w_twips = int(elem.get("width-twips", "2000"))
    h_twips = int(elem.get("height-twips", "1000"))
    has_border = parse_bool(elem.get("has-border"))
    border_color = elem.get("border-color", "000000")

    paragraphs = elem.findall("paragraph")
    if not paragraphs:
        # If no paragraph children, create one from direct text
        para = doc.add_paragraph()
        text = (elem.text or "").strip()
        if text:
            run = para.add_run(text)
            run.font.name = page_font_latin
            run.font.size = Pt(10)
        add_frame_pr(para, x_twips, y_twips, w_twips, h_twips)
        if has_border:
            add_paragraph_borders(para, border_color)
        return

    for p_elem in paragraphs:
        para = doc.add_paragraph()
        alignment = get_alignment(p_elem.get("alignment"))
        if alignment is not None:
            para.alignment = alignment
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)

        process_runs(para, p_elem, page_font_latin, page_font_cjk)
        add_frame_pr(para, x_twips, y_twips, w_twips, h_twips)
        if has_border:
            add_paragraph_borders(para, border_color)


def process_side_by_side(doc, elem, page_font_latin, page_font_cjk):
    cols_count = int(elem.get("cols", "2"))
    columns = elem.findall("column")
    if not columns:
        return

    # Find max rows across columns
    max_rows = 0
    col_paragraphs = []
    for col_elem in columns:
        paras = col_elem.findall("paragraph")
        col_paragraphs.append(paras)
        max_rows = max(max_rows, len(paras))

    if max_rows == 0:
        return

    table = doc.add_table(rows=max_rows, cols=cols_count)
    table.autofit = False

    # Set invisible borders
    for row in table.rows:
        for cell in row.cells:
            set_invisible_cell_borders(cell)

    for col_idx, paras in enumerate(col_paragraphs):
        if col_idx >= cols_count:
            break
        for row_idx, p_elem in enumerate(paras):
            if row_idx >= max_rows:
                break
            cell = table.cell(row_idx, col_idx)
            cell.paragraphs[0].clear()
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)

            alignment = get_alignment(p_elem.get("alignment"))
            if alignment is not None:
                para.alignment = alignment

            process_runs(para, p_elem, page_font_latin, page_font_cjk)

    return table


# ---------------------------------------------------------------------------
# Main page processor
# ---------------------------------------------------------------------------


def process_page(doc, page_xml_path, workspace, is_first_page):
    tree = etree.parse(str(page_xml_path))
    root = tree.getroot()

    if root.tag != "page":
        print(
            f"Warning: unexpected root tag '{root.tag}' in {page_xml_path}",
            file=sys.stderr,
        )
        return

    # Read page attributes
    page_font_latin = root.get("font-latin", "Arial")
    page_font_cjk = root.get("font-cjk", "SimSun")

    if not is_first_page:
        add_page_break(doc)

    # Process child elements in order
    for child in root:
        tag = child.tag
        if tag == "heading":
            process_heading(doc, child, page_font_latin, page_font_cjk)
        elif tag == "paragraph":
            process_paragraph(doc, child, page_font_latin, page_font_cjk)
        elif tag == "table":
            process_table(doc, child, page_font_latin, page_font_cjk, workspace)
        elif tag == "image":
            process_image(doc, child, workspace, page_font_latin, page_font_cjk)
        elif tag == "text-frame":
            process_text_frame(doc, child, page_font_latin, page_font_cjk)
        elif tag == "side-by-side":
            process_side_by_side(doc, child, page_font_latin, page_font_cjk)
        elif tag == "col-widths":
            pass  # Handled within table processing
        else:
            print(
                f"Warning: unknown element <{tag}> in {page_xml_path}", file=sys.stderr
            )


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main():
    parser = argparse.ArgumentParser(description="Convert XML DSL pages to DOCX")
    parser.add_argument("--workspace", required=True, help="Workspace directory path")
    parser.add_argument("--output", required=True, help="Output DOCX file path")
    parser.add_argument(
        "--dsl-dir",
        default="dsl",
        help="DSL directory name under workspace (default: dsl)",
    )
    args = parser.parse_args()

    workspace = Path(args.workspace)
    dsl_dir = workspace / args.dsl_dir
    output_path = Path(args.output)

    # Find and sort page XML files
    xml_files = sorted(
        dsl_dir.glob("page-*.xml"),
        key=lambda p: int(re.search(r"page-(\d+)", p.stem).group(1)),
    )

    if not xml_files:
        print(f"Error: no page-*.xml files found in {dsl_dir}", file=sys.stderr)
        sys.exit(1)

    print(f"Found {len(xml_files)} page XML files in {dsl_dir}")

    # Read page setup from first page XML
    first_tree = etree.parse(str(xml_files[0]))
    first_root = first_tree.getroot()

    width_pts = float(first_root.get("width-pts", "595.276"))
    height_pts = float(first_root.get("height-pts", "841.89"))
    margin_top = float(first_root.get("margin-top-cm", "1.27"))
    margin_bottom = float(first_root.get("margin-bottom-cm", "1.27"))
    margin_left = float(first_root.get("margin-left-cm", "1.27"))
    margin_right = float(first_root.get("margin-right-cm", "1.27"))

    # Create document
    doc = Document()
    section = doc.sections[0]
    section.page_width = Emu(pts_to_emu(width_pts))
    section.page_height = Emu(pts_to_emu(height_pts))
    section.top_margin = Cm(margin_top)
    section.bottom_margin = Cm(margin_bottom)
    section.left_margin = Cm(margin_left)
    section.right_margin = Cm(margin_right)

    # Process each page
    for i, xml_file in enumerate(xml_files):
        page_num = int(re.search(r"page-(\d+)", xml_file.stem).group(1))
        print(f"Processing page {page_num} ({xml_file.name})")
        process_page(doc, xml_file, workspace, is_first_page=(i == 0))

    # Save
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Saved DOCX to {output_path}")


if __name__ == "__main__":
    main()
