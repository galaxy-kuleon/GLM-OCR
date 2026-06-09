#!/usr/bin/env python3
"""Run GLM-OCR → DocIR → DOCX battle tests on selected PDFs.

This harness is intentionally conservative:
- digital PDFs use --digital fast path
- non-digital PDFs skip VLM style extraction by default for survival smoke tests
- every case has an independent work dir, log, and result JSON
"""
from __future__ import annotations

import argparse
import json
import subprocess
import sys
import time
from datetime import datetime
from pathlib import Path
from typing import Any

try:
    from docx import Document
except Exception:  # pragma: no cover
    Document = None

try:
    import lxml.etree as etree
except Exception:  # pragma: no cover
    etree = None

REPO = Path(__file__).resolve().parents[2]
IR_SCHEMA = REPO / "ir-schema"
SELECTED = IR_SCHEMA / "battle_tests" / "selected_pdfs.json"


def slugify(rel: str) -> str:
    s = rel.replace("/", "__")
    keep = []
    for ch in s:
        if ch.isalnum() or ch in "._-":
            keep.append(ch)
        else:
            keep.append("_")
    out = "".join(keep).strip("_")
    return out[:160]


def docx_metrics(path: Path) -> dict[str, Any]:
    if not path.exists() or Document is None:
        return {}
    try:
        doc = Document(str(path))
        return {
            "sections": len(doc.sections),
            "paragraphs": len(doc.paragraphs),
            "tables": len(doc.tables),
            "body_chars": sum(len(p.text or "") for p in doc.paragraphs),
        }
    except Exception as e:
        return {"error": str(e)}


def docir_metrics(work_dir: Path, stem: str) -> dict[str, Any]:
    if etree is None:
        return {}
    candidates = list(work_dir.rglob(f"{stem}*docir.xml"))
    if not candidates:
        return {}
    # Prefer styled final docir if present.
    candidates.sort(key=lambda p: ("styled" not in p.name, len(p.name)))
    path = candidates[0]
    try:
        root = etree.parse(str(path)).getroot()
        ns = {"d": "urn:docir:v0.1"}
        return {
            "docir": str(path),
            "pages": len(root.findall(".//d:page", ns)),
            "regions": len(root.findall(".//d:region", ns)),
            "tables": len(root.findall(".//d:region[@type='table']", ns)),
            "images": len(root.findall(".//d:region[@type='image']", ns)),
            "headers": len(root.findall(".//d:region[@role='header']", ns)),
            "footers": len(root.findall(".//d:region[@role='footer']", ns)),
        }
    except Exception as e:
        return {"docir": str(path), "error": str(e)}


def prefer_digital_direct(case: dict[str, Any]) -> bool:
    """Return True when a case should use PyMuPDF digital-direct before OCR.

    Human-POV lesson from battle testing: low-text native PDFs (forms, tables,
    hidden OCR layers, mixed native+raster) can be misclassified as
    "mixed/low-text" and then spend 4+ minutes timing out in Ollama OCR. If the
    PDF exposes any native text or image structure, digital-direct gives an
    immediate editable-or-at-least-visible DOCX candidate; visual evaluation can
    then decide whether hidden text / screenshot-like output is unacceptable.
    """
    if case.get("kind") == "digital":
        return True
    if int(case.get("avg_text") or 0) > 0:
        return True
    if int(case.get("images") or 0) > 0 and case.get("kind") != "scanned/image":
        return True
    tags = set(case.get("tags") or [])
    if "table" in tags or "adversarial" in tags:
        return int(case.get("avg_text") or 0) > 0
    return False


def run_case(case: dict[str, Any], run_dir: Path, timeout: int, deep_style: bool) -> dict[str, Any]:
    pdf = Path(case["path"])
    rel = case["rel"]
    slug = slugify(rel)
    case_dir = run_dir / slug
    work_dir = case_dir / "work"
    case_dir.mkdir(parents=True, exist_ok=True)
    work_dir.mkdir(parents=True, exist_ok=True)
    out_docx = case_dir / f"{pdf.stem}.docx"
    log_path = case_dir / "pipeline.log"

    cmd = [
        sys.executable,
        str(IR_SCHEMA / "run_pipeline.py"),
        str(pdf),
        "-o", str(out_docx),
        "--work-dir", str(work_dir),
        "--positioned",
        "--detect-headers",
    ]

    # Digital-direct first for any PDF with native extractable structure. This is not
    # a quality pass; it is a fast candidate generation strategy. Human-POV visual
    # review remains the gate for whether the output is actually useful.
    use_digital = prefer_digital_direct(case)
    if use_digital:
        cmd.append("--digital")
    elif not deep_style:
        cmd.append("--skip-style")
    else:
        cmd.extend(["--parallel", "2"])

    started = time.time()
    result: dict[str, Any] = {
        "rel": rel,
        "path": str(pdf),
        "kind": case.get("kind"),
        "strategy": "digital-direct" if use_digital else ("ocr-vlm-style" if deep_style else "ocr-skip-style"),
        "tags": case.get("tags", []),
        "pages_expected": case.get("pages"),
        "size_mb": case.get("size_mb"),
        "cmd": cmd,
        "started_at": datetime.now().isoformat(timespec="seconds"),
    }

    with log_path.open("w") as log:
        log.write("$ " + " ".join(cmd) + "\n\n")
        log.flush()
        try:
            proc = subprocess.run(
                cmd,
                cwd=str(REPO),
                stdout=log,
                stderr=subprocess.STDOUT,
                timeout=timeout,
            )
            result["exit_code"] = proc.returncode
            result["status"] = "passed" if proc.returncode == 0 and out_docx.exists() else "failed"
        except subprocess.TimeoutExpired:
            result["exit_code"] = None
            result["status"] = "timeout"
            log.write(f"\nTIMEOUT after {timeout}s\n")
        except Exception as e:
            result["exit_code"] = None
            result["status"] = "error"
            result["error"] = str(e)

    result["duration_s"] = round(time.time() - started, 1)
    result["output_docx"] = str(out_docx)
    result["log"] = str(log_path)
    result["docx_metrics"] = docx_metrics(out_docx)
    result["docir_metrics"] = docir_metrics(work_dir, pdf.stem)

    (case_dir / "result.json").write_text(json.dumps(result, indent=2, ensure_ascii=False))
    return result


def write_summary(results: list[dict[str, Any]], run_dir: Path) -> None:
    lines = [
        "# Battle Test Summary",
        "",
        f"Run dir: `{run_dir}`",
        f"Generated: {datetime.now().isoformat(timespec='seconds')}",
        "",
        "| # | Status | Time | Kind | Strategy | Pages | Regions | Tables | DOCX | Case |",
        "|---:|---|---:|---|---|---:|---:|---:|---|---|",
    ]
    for i, r in enumerate(results, 1):
        dm = r.get("docir_metrics") or {}
        dx = r.get("docx_metrics") or {}
        docx_cell = "yes" if Path(r.get("output_docx", "")).exists() else "no"
        lines.append(
            f"| {i} | {r.get('status')} | {r.get('duration_s')}s | {r.get('kind')} | {r.get('strategy','')} | "
            f"{r.get('pages_expected')} | {dm.get('regions','')} | {dm.get('tables','')} | "
            f"{docx_cell} ({dx.get('paragraphs','')}p/{dx.get('tables','')}t) | `{r.get('rel')}` |"
        )
    lines.append("")
    failed = [r for r in results if r.get("status") != "passed"]
    if failed:
        lines.append("## Failures / Timeouts")
        for r in failed:
            lines.append(f"- `{r.get('rel')}` → {r.get('status')} ({r.get('log')})")
    else:
        lines.append("All cases passed survival pipeline.")
    (run_dir / "summary.md").write_text("\n".join(lines), encoding="utf-8")
    (run_dir / "results.json").write_text(json.dumps(results, indent=2, ensure_ascii=False), encoding="utf-8")


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--selected", type=Path, default=SELECTED)
    ap.add_argument("--run-dir", type=Path, default=None)
    ap.add_argument("--limit", type=int, default=0)
    ap.add_argument("--indices", help="Comma-separated 1-based selected_pdfs indices to run")
    ap.add_argument("--timeout", type=int, default=360)
    ap.add_argument("--deep-style", action="store_true", help="Use VLM style extraction for non-digital PDFs")
    args = ap.parse_args()

    cases = json.loads(args.selected.read_text())
    if args.indices:
        wanted = {int(x.strip()) for x in args.indices.split(",") if x.strip()}
        cases = [c for i, c in enumerate(cases, 1) if i in wanted]
    if args.limit:
        cases = cases[:args.limit]

    if args.run_dir is None:
        args.run_dir = IR_SCHEMA / "battle_tests" / "runs" / datetime.now().strftime("%Y%m%d-%H%M%S")
    args.run_dir.mkdir(parents=True, exist_ok=True)

    print(f"Running {len(cases)} cases → {args.run_dir}")
    results = []
    for i, case in enumerate(cases, 1):
        print(f"[{i}/{len(cases)}] {case['rel']} ({case.get('kind')}, {case.get('pages')}p)")
        r = run_case(case, args.run_dir, args.timeout, args.deep_style)
        results.append(r)
        print(f"  → {r['status']} in {r['duration_s']}s; log={r['log']}")
        write_summary(results, args.run_dir)

    write_summary(results, args.run_dir)
    print(f"\nSummary: {args.run_dir / 'summary.md'}")
    return 0 if all(r.get("status") == "passed" for r in results) else 1


if __name__ == "__main__":
    raise SystemExit(main())
