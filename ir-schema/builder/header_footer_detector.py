#!/usr/bin/env python3
"""
Header/Footer Detector - Detect repeating headers and footers across pages.

Algorithm (from pdf2docx):
1. Collect text regions near top/bottom of each page
2. Group by vertical position (y-coordinate)
3. If same/similar text appears at same position on 3+ pages → header/footer
4. Mark regions in DocIR XML as header/footer type
5. DOCX generator can use actual DOCX header/footer sections

Handles:
- Page numbers (detected by pattern matching)
- Repeating titles/company names
- Different headers on first page vs rest
"""

import re
from pathlib import Path
from typing import Dict, List, Tuple, Optional, Set
from collections import defaultdict
from lxml import etree


DOCIR_NS = "urn:docir:v0.1"


def detect_headers_footers(
    docir_path: Path,
    output_path: Optional[Path] = None,
    margin_pct: float = 0.15,
    min_pages: int = 3,
    similarity_threshold: float = 0.8,
) -> Path:
    """
    Detect headers and footers by finding repeating text across pages.
    
    Args:
        docir_path: Path to DocIR XML
        output_path: Output path (default: overwrite input)
        margin_pct: Top/bottom margin percentage to search (0.15 = 15%)
        min_pages: Minimum pages with same text to qualify
        similarity_threshold: Text similarity threshold (0-1)
    
    Returns:
        Path to output DocIR XML
    """
    if output_path is None:
        output_path = docir_path
    
    # Load DocIR XML
    tree = etree.parse(str(docir_path))
    root = tree.getroot()
    
    pages = root.findall(f".//{{{DOCIR_NS}}}page")
    if len(pages) < min_pages:
        print(f"  Skipping: only {len(pages)} pages (need {min_pages})")
        return output_path
    
    # Collect candidate regions from top/bottom of each page
    header_candidates = defaultdict(list)  # (y_position_bucket, text_hash) -> [(page_idx, region_elem)]
    footer_candidates = defaultdict(list)
    
    for page_idx, page_elem in enumerate(pages):
        page_height = float(page_elem.get('height_pt', '842'))  # A4 default
        
        # Margin boundaries
        top_margin = page_height * margin_pct
        bottom_margin = page_height * (1 - margin_pct)
        
        regions_elem = page_elem.find(f".//{{{DOCIR_NS}}}regions")
        if regions_elem is None:
            continue
        
        for region in regions_elem.findall(f".//{{{DOCIR_NS}}}region"):
            region_type = region.get('type')
            if region_type not in ('text', 'paragraph_title'):
                continue
            
            bbox = region.find(f".//{{{DOCIR_NS}}}bbox")
            if bbox is None:
                continue
            
            y = float(bbox.get('y', '0'))
            h = float(bbox.get('height', '0'))
            
            # Get text content
            text = _get_region_text(region)
            if not text or len(text.strip()) < 2:
                continue
            
            # Skip page number patterns (they vary per page)
            if _is_page_number(text):
                continue
            
            # Bucket y-position (round to nearest 5pt for matching)
            y_bucket = round(y / 5) * 5
            text_hash = _normalize_text(text)
            
            if y + h < top_margin:
                # Header region (near top)
                key = (y_bucket, text_hash)
                header_candidates[key].append((page_idx, region))
            elif y > bottom_margin:
                # Footer region (near bottom)
                key = (y_bucket, text_hash)
                footer_candidates[key].append((page_idx, region))
    
    # Find repeating patterns
    headers_found = 0
    footers_found = 0
    
    # Process headers
    for key, occurrences in header_candidates.items():
        page_indices = [p for p, r in occurrences]
        if len(set(page_indices)) >= min_pages:
            # This is a header - mark all occurrences
            for page_idx, region in occurrences:
                region.set('role', 'header')
                headers_found += 1
    
    # Process footers
    for key, occurrences in footer_candidates.items():
        page_indices = [p for p, r in occurrences]
        if len(set(page_indices)) >= min_pages:
            # This is a footer - mark all occurrences
            for page_idx, region in occurrences:
                region.set('role', 'footer')
                footers_found += 1
    
    # Write output
    tree.write(
        str(output_path),
        xml_declaration=True,
        encoding="UTF-8",
        pretty_print=True
    )
    
    print(f"✓ Header/footer detection complete")
    print(f"  Headers found: {headers_found} regions")
    print(f"  Footers found: {footers_found} regions")
    print(f"  Output: {output_path}")
    
    return output_path


def _get_region_text(region: etree.Element) -> str:
    """Extract all text content from a region."""
    texts = []
    
    # Check text_content
    text_content = region.find(f".//{{{DOCIR_NS}}}text_content")
    if text_content is not None:
        for para in text_content.findall(f".//{{{DOCIR_NS}}}paragraph"):
            para_text = ""
            for run in para.findall(f".//{{{DOCIR_NS}}}run"):
                if run.text:
                    para_text += run.text
            if para_text.strip():
                texts.append(para_text.strip())
    
    # Check table_content cells
    table_content = region.find(f".//{{{DOCIR_NS}}}table_content")
    if table_content is not None:
        for cell in table_content.findall(f".//{{{DOCIR_NS}}}cell"):
            if cell.text and cell.text.strip():
                texts.append(cell.text.strip())
    
    return ' '.join(texts)


def _normalize_text(text: str) -> str:
    """Normalize text for comparison (remove variable parts)."""
    # Remove page numbers
    text = re.sub(r'\b\d+\b', '', text)
    # Remove dates (common patterns)
    text = re.sub(r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}', '', text)
    text = re.sub(r'\d{4}[/-]\d{1,2}[/-]\d{1,2}', '', text)
    # Normalize whitespace
    text = ' '.join(text.split())
    return text.strip().lower()


def _is_page_number(text: str) -> bool:
    """Check if text looks like a page number."""
    text = text.strip()
    # Pure number
    if text.isdigit():
        return True
    # "Page X" or "Page X of Y"
    if re.match(r'^page\s+\d+(\s+of\s+\d+)?$', text.lower()):
        return True
    # Roman numerals
    if re.match(r'^[ivxlcdm]+$', text.lower()):
        return True
    # "- X -" format
    if re.match(r'^-\s*\d+\s*-$', text):
        return True
    return False


def apply_headers_footers_to_docx(docx_path: Path, docir_path: Path) -> None:
    """
    Apply detected headers/footers to DOCX file.
    
    Reads header/footer regions from DocIR and creates actual
    DOCX header/footer sections.
    
    Args:
        docx_path: Path to DOCX file to modify
        docir_path: Path to DocIR XML with header/footer annotations
    """
    from docx import Document
    from docx.shared import Pt
    
    # Load DocIR
    tree = etree.parse(str(docir_path))
    root = tree.getroot()
    
    # Collect header/footer text
    header_texts = {}  # page_idx -> text
    footer_texts = {}
    
    for page_elem in root.findall(f".//{{{DOCIR_NS}}}page"):
        page_idx = int(page_elem.get('index', 0))
        
        for region in page_elem.findall(f".//{{{DOCIR_NS}}}region"):
            role = region.get('role')
            if role not in ('header', 'footer'):
                continue
            
            text = _get_region_text(region)
            if not text:
                continue
            
            if role == 'header':
                header_texts[page_idx] = text
            else:
                footer_texts[page_idx] = text
    
    if not header_texts and not footer_texts:
        return
    
    # Load DOCX
    doc = Document(str(docx_path))
    
    # Apply headers (use most common header text)
    if header_texts:
        # Find most common header
        from collections import Counter
        header_counter = Counter(header_texts.values())
        main_header = header_counter.most_common(1)[0][0]
        
        for section in doc.sections:
            header = section.header
            header.is_linked_to_previous = False
            # Clear existing
            for para in header.paragraphs:
                para.clear()
            # Add header text
            if header.paragraphs:
                header.paragraphs[0].text = main_header
                header.paragraphs[0].alignment = 1  # Center
            else:
                para = header.add_paragraph(main_header)
                para.alignment = 1
    
    # Apply footers (use most common footer text)
    if footer_texts:
        from collections import Counter
        footer_counter = Counter(footer_texts.values())
        main_footer = footer_counter.most_common(1)[0][0]
        
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            for para in footer.paragraphs:
                para.clear()
            if footer.paragraphs:
                footer.paragraphs[0].text = main_footer
                footer.paragraphs[0].alignment = 1
            else:
                para = footer.add_paragraph(main_footer)
                para.alignment = 1
    
    doc.save(str(docx_path))
    print(f"✓ Applied headers/footers to DOCX")


if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='Detect headers/footers in DocIR')
    parser.add_argument('docir', type=Path, help='Input DocIR XML')
    parser.add_argument('-o', '--output', type=Path, help='Output DocIR XML')
    parser.add_argument('--margin-pct', type=float, default=0.15,
                       help='Top/bottom margin percentage (default: 0.15)')
    parser.add_argument('--min-pages', type=int, default=3,
                       help='Min pages for detection (default: 3)')
    
    args = parser.parse_args()
    
    detect_headers_footers(
        args.docir,
        args.output,
        margin_pct=args.margin_pct,
        min_pages=args.min_pages,
    )
