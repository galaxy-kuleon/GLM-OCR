#!/usr/bin/env python3
"""
DocIR to DOCX Generator v1.1

Converts DocIR XML to editable DOCX using python-docx.
Preserves text content and typography styles (font, size, bold, italic, color).

Modes:
  - Default (flow): Linear document flow, regions become paragraphs
  - Positioned (--positioned): Absolute positioning using page coordinates

Usage:
    python docx_generator.py input.docir.xml -o output.docx
    python docx_generator.py input.docir.xml -o output.docx --positioned
"""

import sys
from pathlib import Path
from typing import Optional, Tuple
from lxml import etree
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# DocIR namespace
DOCIR_NS = "urn:docir:v0.1"

# Conversion constants
PT_TO_EMU = 12700  # 1 point = 12700 EMU


def hex_to_rgb(hex_color: str) -> Optional[RGBColor]:
    """Convert hex color (#RRGGBB) to RGBColor."""
    if not hex_color or not hex_color.startswith('#'):
        return None
    
    hex_color = hex_color.lstrip('#')
    if len(hex_color) != 6:
        return None
    
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return RGBColor(r, g, b)
    except ValueError:
        return None


def set_char_scaling(run, scale: float):
    """
    Set character scaling (horizontal stretch/compress).
    
    Args:
        run: python-docx Run object
        scale: Scaling factor (1.0 = 100% = normal, 0.5 = 50%, 2.0 = 200%)
    """
    # w:w value is in percent (100 = normal)
    scale_percent = int(100 * scale)
    xml = f'<w:w {nsdecls("w")} w:val="{scale_percent}"/>'
    run._r.get_or_add_rPr().insert(0, parse_xml(xml))


def set_char_spacing(run, spacing_pt: float):
    """
    Set character spacing (kerning).
    
    Args:
        run: python-docx Run object
        spacing_pt: Spacing in points (positive = expand, negative = condense)
    """
    # w:spacing value is in half-points (20 twips = 1 point)
    spacing_twips = int(20 * spacing_pt)
    xml = f'<w:spacing {nsdecls("w")} w:val="{spacing_twips}"/>'
    run._r.get_or_add_rPr().insert(0, parse_xml(xml))


def set_char_shading(run, color_hex: str):
    """
    Set character background shading.
    
    Args:
        run: python-docx Run object
        color_hex: Hex color string (#RRGGBB)
    """
    if not color_hex or not color_hex.startswith('#'):
        return
    
    color = color_hex.lstrip('#')
    if len(color) != 6:
        return
    
    xml = f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color}"/>'
    run._r.get_or_add_rPr().insert(0, parse_xml(xml))


def set_underline_color(run, color_hex: str):
    """
    Set underline with specific color.
    
    Args:
        run: python-docx Run object
        color_hex: Hex color string (#RRGGBB)
    """
    if not color_hex or not color_hex.startswith('#'):
        return
    
    color = color_hex.lstrip('#')
    if len(color) != 6:
        return
    
    xml = f'<w:u {nsdecls("w")} w:val="single" w:color="{color}"/>'
    run._r.get_or_add_rPr().insert(0, parse_xml(xml))


def apply_run_style(run, run_elem, region_elem=None):
    """Apply style attributes from DocIR run element to python-docx run.
    
    Priority:
    1. Run-level attributes (font_name, font_size_pt, etc.)
    2. Region-level computed_style (from digital_extractor or style_extractor)
    """
    font = run.font
    
    # First, try to get styles from region's computed_style
    computed_style = None
    if region_elem is not None:
        computed_style = region_elem.find(f'{{{DOCIR_NS}}}computed_style')
    
    # Font name (run attr > computed_style > default)
    font_name = run_elem.get('font_name')
    if not font_name and computed_style is not None:
        font_elem = computed_style.find(f'{{{DOCIR_NS}}}font')
        if font_elem is not None:
            font_name = font_elem.get('family') or font_elem.get('name')
    if font_name:
        font.name = font_name
    
    # Font size (run attr > computed_style > default)
    font_size_pt = run_elem.get('font_size_pt')
    if not font_size_pt and computed_style is not None:
        font_elem = computed_style.find(f'{{{DOCIR_NS}}}font')
        if font_elem is not None:
            font_size_pt = font_elem.get('size_pt')
    if font_size_pt:
        try:
            font.size = Pt(float(font_size_pt))
        except ValueError:
            pass
    
    # Bold (run attr > computed_style > default)
    bold = run_elem.get('bold')
    if not bold and computed_style is not None:
        font_elem = computed_style.find(f'{{{DOCIR_NS}}}font')
        if font_elem is not None:
            bold = font_elem.get('bold')
    if bold:
        font.bold = bold.lower() == 'true'
    
    # Italic (run attr > computed_style > default)
    italic = run_elem.get('italic')
    if not italic and computed_style is not None:
        font_elem = computed_style.find(f'{{{DOCIR_NS}}}font')
        if font_elem is not None:
            italic = font_elem.get('italic')
    if italic:
        font.italic = italic.lower() == 'true'
    
    # Underline (run attr only for now)
    underline = run_elem.get('underline')
    if underline:
        font.underline = underline.lower() == 'true'
    
    # Color (run attr > computed_style > default)
    color = run_elem.get('color')
    if not color and computed_style is not None:
        color_elem = computed_style.find(f'{{{DOCIR_NS}}}color')
        if color_elem is not None:
            color = color_elem.get('value')
    if color:
        rgb = hex_to_rgb(color)
        if rgb:
            font.color.rgb = rgb
    
    # Rich formatting properties (from computed_style only)
    if computed_style is not None:
        # Character scaling (horizontal stretch)
        scale_elem = computed_style.find(f'{{{DOCIR_NS}}}text_scale')
        if scale_elem is not None:
            scale_val = scale_elem.get('percent')
            if scale_val:
                try:
                    scale_factor = float(scale_val) / 100.0
                    set_char_scaling(run, scale_factor)
                except ValueError:
                    pass
        
        # Character spacing (kerning)
        spacing_elem = computed_style.find(f'{{{DOCIR_NS}}}char_spacing')
        if spacing_elem is not None:
            spacing_pt = spacing_elem.get('pt')
            if spacing_pt:
                try:
                    set_char_spacing(run, float(spacing_pt))
                except ValueError:
                    pass
        
        # Character shading (background color)
        shading_elem = computed_style.find(f'{{{DOCIR_NS}}}shading')
        if shading_elem is not None:
            shading_color = shading_elem.get('fill')
            if shading_color:
                set_char_shading(run, shading_color)
        
        # Underline with color
        underline_color_elem = computed_style.find(f'{{{DOCIR_NS}}}underline_color')
        if underline_color_elem is not None:
            ul_color = underline_color_elem.get('value')
            if ul_color:
                set_underline_color(run, ul_color)


def set_paragraph_alignment(paragraph, alignment_str: str):
    """Set paragraph alignment from DocIR alignment attribute."""
    if not alignment_str:
        return
    
    alignment_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    
    alignment = alignment_map.get(alignment_str.lower())
    if alignment:
        paragraph.alignment = alignment


def get_bbox(region_elem) -> Optional[Tuple[float, float, float, float]]:
    """Extract bbox (x, y, width, height) from region element in PDF points."""
    bbox_elem = region_elem.find(f'{{{DOCIR_NS}}}bbox')
    if bbox_elem is None:
        return None
    
    try:
        x = float(bbox_elem.get('x', 0))
        y = float(bbox_elem.get('y', 0))
        w = float(bbox_elem.get('width', 0))
        h = float(bbox_elem.get('height', 0))
        return (x, y, w, h)
    except ValueError:
        return None


def setup_positioned_section(section, page_width_pt: float, page_height_pt: float):
    """Configure section for absolute positioning (zero margins, correct page size)."""
    # Set page dimensions (convert pt to inches for python-docx)
    section.page_width = Inches(page_width_pt / 72.0)
    section.page_height = Inches(page_height_pt / 72.0)
    
    # Set margins to 0 for absolute positioning
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)
    
    # Set orientation
    if page_width_pt > page_height_pt:
        section.orientation = WD_ORIENT.LANDSCAPE
    else:
        section.orientation = WD_ORIENT.PORTRAIT


def add_positioned_text_region(doc: Document, region_elem, page_height_pt: float, current_y_pt: float) -> float:
    """
    Add a text region with absolute positioning.
    
    Returns the new current_y_pt (bottom of this region).
    """
    # Check if this is a floating/text box region
    is_floating = region_elem.get('floating') == 'true'
    if is_floating:
        return add_floating_text_box(doc, region_elem, page_height_pt, current_y_pt)
    
    bbox = get_bbox(region_elem)
    if bbox is None:
        return current_y_pt
    
    x_pt, y_pt, w_pt, h_pt = bbox
    
    # Convert DocIR coordinates to page-top-relative coordinates
    # DocIR: y is from bottom, so top_of_region_from_page_top = page_height - y
    top_from_page_top_pt = page_height_pt - y_pt
    left_pt = x_pt
    
    # Calculate space needed before this region
    space_before_pt = top_from_page_top_pt - current_y_pt
    if space_before_pt < 0:
        space_before_pt = 0
    
    text_content = region_elem.find(f'{{{DOCIR_NS}}}text_content')
    if text_content is None:
        return current_y_pt
    
    # Process each paragraph in the region
    for para_elem in text_content.findall(f'{{{DOCIR_NS}}}paragraph'):
        para = doc.add_paragraph()
        
        # Set vertical position via space_before
        pf = para.paragraph_format
        pf.space_before = Pt(space_before_pt)
        pf.space_after = Pt(0)
        
        # Set horizontal position via left_indent
        pf.left_indent = Pt(left_pt)
        
        # Paragraph alignment (para attr > region computed_style > default)
        alignment = para_elem.get('alignment')
        if not alignment and region_elem is not None:
            computed_style = region_elem.find(f'{{{DOCIR_NS}}}computed_style')
            if computed_style is not None:
                align_elem = computed_style.find(f'{{{DOCIR_NS}}}alignment')
                if align_elem is not None:
                    alignment = align_elem.get('value')
        if alignment:
            set_paragraph_alignment(para, alignment)
        
        # Process each run
        for run_elem in para_elem.findall(f'{{{DOCIR_NS}}}run'):
            text = run_elem.text or ''
            if text:
                run = para.add_run(text)
                apply_run_style(run, run_elem, region_elem)
        
        # After first paragraph, reset space_before for subsequent paragraphs
        space_before_pt = 0
    
    # Return the bottom of this region
    return top_from_page_top_pt + h_pt


def add_floating_text_box(doc: Document, region_elem, page_height_pt: float, current_y_pt: float) -> float:
    """
    Add a floating text box region (aside_text, header, footer).
    
    Renders as a single-cell bordered table to simulate a text box.
    The table has:
    - A thin border around the cell
    - The text content inside
    - Width matching the bbox
    - Background shading (light gray) to distinguish from main text
    """
    bbox = get_bbox(region_elem)
    if bbox is None:
        return current_y_pt
    
    x_pt, y_pt, w_pt, h_pt = bbox
    top_from_page_top_pt = page_height_pt - y_pt
    
    # Calculate space needed before this region
    space_before_pt = top_from_page_top_pt - current_y_pt
    if space_before_pt < 0:
        space_before_pt = 0
    
    text_content = region_elem.find(f'{{{DOCIR_NS}}}text_content')
    if text_content is None:
        return current_y_pt
    
    # Create a single-cell table to simulate text box
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    
    # Set table width to match bbox
    table.columns[0].width = Pt(w_pt)
    
    # Get the cell
    cell = table.cell(0, 0)
    
    # Add spacing before the table
    # We add an empty paragraph before the table for vertical positioning
    table._element.getparent().insert(
        table._element.getparent().index(table._element),
        doc.add_paragraph()._element
    )
    spacer = doc.paragraphs[-2] if len(doc.paragraphs) > 1 else doc.paragraphs[0]
    spacer.paragraph_format.space_before = Pt(space_before_pt)
    spacer.paragraph_format.space_after = Pt(0)
    
    # Set cell borders (thin black border)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    # Add borders
    tcBorders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # 0.5pt
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '666666')
        tcBorders.append(border)
    tcPr.append(tcBorders)
    
    # Add light gray shading
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), 'F5F5F5')
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)
    
    # Set cell width
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(w_pt * 20)))  # Convert pt to twips
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    
    # Clear default paragraph and add content
    cell.paragraphs[0].clear()
    
    first_para = True
    for para_elem in text_content.findall(f'{{{DOCIR_NS}}}paragraph'):
        if first_para:
            para = cell.paragraphs[0]
            first_para = False
        else:
            para = cell.add_paragraph()
        
        # Process each run
        for run_elem in para_elem.findall(f'{{{DOCIR_NS}}}run'):
            text = run_elem.text or ''
            if text:
                run = para.add_run(text)
                apply_run_style(run, run_elem, region_elem)
    
    # Return the bottom of this region
    return top_from_page_top_pt + h_pt


def add_positioned_table_region(doc: Document, region_elem, page_height_pt: float, current_y_pt: float) -> float:
    """
    Add a table region with absolute positioning.
    
    Returns the new current_y_pt (bottom of this region).
    """
    bbox = get_bbox(region_elem)
    if bbox is None:
        return current_y_pt
    
    x_pt, y_pt, w_pt, h_pt = bbox
    top_from_page_top_pt = page_height_pt - y_pt
    left_pt = x_pt
    
    # Calculate space needed before this region
    space_before_pt = top_from_page_top_pt - current_y_pt
    if space_before_pt < 0:
        space_before_pt = 0
    
    # Add spacer paragraph for vertical positioning
    if space_before_pt > 0:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(space_before_pt)
        spacer.paragraph_format.space_after = Pt(0)
    
    # Process table content
    table_content = region_elem.find(f'{{{DOCIR_NS}}}table_content')
    if table_content is None:
        return current_y_pt
    
    num_rows = int(table_content.get('rows', '0'))
    num_cols = int(table_content.get('cols', '0'))
    
    if num_rows == 0 or num_cols == 0:
        return current_y_pt
    
    # Get table style
    table_style_elem = table_content.find(f'{{{DOCIR_NS}}}table_style')
    border_visible = True
    border_color = "000000"
    has_header = False
    
    if table_style_elem is not None:
        border_visible = table_style_elem.get('border_visible', 'true').lower() == 'true'
        border_color_hex = table_style_elem.get('border_color', '#000000')
        border_color = border_color_hex.lstrip('#')
        has_header = table_style_elem.get('header_row', 'false').lower() == 'true'
    
    # Create table
    table = doc.add_table(rows=num_rows, cols=num_cols)
    
    # Set table position via paragraph format on the table's containing paragraph
    tbl_para = table._tbl.getparent()
    if hasattr(tbl_para, 'paragraph_format'):
        pass  # Table is already placed
    
    # Apply borders if visible
    if border_visible:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)
            borders.append(border)
        
        tblPr.append(borders)
    
    # Helper: apply per-cell borders from cell_borders element
    def apply_cell_borders(cell_elem, cell):
        """Apply per-cell border styling from DocIR cell_borders element."""
        cell_borders_elem = cell_elem.find(f'{{{DOCIR_NS}}}cell_borders')
        if cell_borders_elem is None:
            return
        
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        tc = cell._tc
        tcPr = tc.tcPr if tc.tcPr is not None else OxmlElement('w:tcPr')
        if tc.tcPr is None:
            tc.insert(0, tcPr)
        
        tcBorders = OxmlElement('w:tcBorders')
        has_borders = False
        
        for edge in ['top', 'bottom', 'left', 'right']:
            border_elem = cell_borders_elem.find(f'{{{DOCIR_NS}}}border_{edge}')
            if border_elem is not None:
                has_borders = True
                border = OxmlElement(f'w:{edge}')
                
                style = border_elem.get('style', 'single')
                border.set(qn('w:val'), style)
                
                # Width in eighths of a point (0.5pt = 4)
                width_pt = float(border_elem.get('width', '0.5'))
                border.set(qn('w:sz'), str(int(width_pt * 8)))
                border.set(qn('w:space'), '0')
                
                color = border_elem.get('color', '#000000').lstrip('#')
                border.set(qn('w:color'), color)
                
                tcBorders.append(border)
        
        if has_borders:
            tcPr.append(tcBorders)
    
    # Set table indent for horizontal positioning
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    tblInd = OxmlElement('w:tblInd')
    tblInd.set(qn('w:w'), str(int(left_pt * 20)))  # twips (1 pt = 20 twips)
    tblInd.set(qn('w:type'), 'dxa')
    tblPr.append(tblInd)
    
    # Fill in cells (same logic as flow mode)
    row_idx = 0
    
    # Process header row group
    header_group = table_content.find(f'{{{DOCIR_NS}}}row_group[@type="header"]')
    if header_group is not None:
        for row_elem in header_group.findall(f'{{{DOCIR_NS}}}row'):
            if row_idx >= num_rows:
                break
            
            col_idx = 0
            for cell_elem in row_elem.findall(f'{{{DOCIR_NS}}}cell'):
                if col_idx >= num_cols:
                    break
                
                cell = table.cell(row_idx, col_idx)
                
                col_span = int(cell_elem.get('col_span', '1'))
                row_span = int(cell_elem.get('row_span', '1'))
                
                if col_span > 1 or row_span > 1:
                    end_row = min(row_idx + row_span - 1, num_rows - 1)
                    end_col = min(col_idx + col_span - 1, num_cols - 1)
                    cell = table.cell(row_idx, col_idx).merge(table.cell(end_row, end_col))
                
                text_content_elem = cell_elem.find(f'{{{DOCIR_NS}}}text_content')
                if text_content_elem is not None:
                    cell.text = ''
                    for para_elem in text_content_elem.findall(f'{{{DOCIR_NS}}}paragraph'):
                        para = cell.add_paragraph()
                        for run_elem in para_elem.findall(f'{{{DOCIR_NS}}}run'):
                            text = run_elem.text or ''
                            run = para.add_run(text)
                            apply_run_style(run, run_elem, region_elem)
                    
                    if len(cell.paragraphs) > 1 and not cell.paragraphs[0].text:
                        p = cell.paragraphs[0]._element
                        p.getparent().remove(p)
                
                if has_header and row_idx == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                
                # Apply per-cell borders if present
                apply_cell_borders(cell_elem, cell)
                
                col_idx += col_span
            
            row_idx += 1
    
    # Process body row group
    body_group = table_content.find(f'{{{DOCIR_NS}}}row_group[@type="body"]')
    if body_group is not None:
        for row_elem in body_group.findall(f'{{{DOCIR_NS}}}row'):
            if row_idx >= num_rows:
                break
            
            col_idx = 0
            for cell_elem in row_elem.findall(f'{{{DOCIR_NS}}}cell'):
                if col_idx >= num_cols:
                    break
                
                cell = table.cell(row_idx, col_idx)
                
                col_span = int(cell_elem.get('col_span', '1'))
                row_span = int(cell_elem.get('row_span', '1'))
                
                if col_span > 1 or row_span > 1:
                    end_row = min(row_idx + row_span - 1, num_rows - 1)
                    end_col = min(col_idx + col_span - 1, num_cols - 1)
                    cell = table.cell(row_idx, col_idx).merge(table.cell(end_row, end_col))
                
                text_content_elem = cell_elem.find(f'{{{DOCIR_NS}}}text_content')
                if text_content_elem is not None:
                    cell.text = ''
                    for para_elem in text_content_elem.findall(f'{{{DOCIR_NS}}}paragraph'):
                        para = cell.add_paragraph()
                        for run_elem in para_elem.findall(f'{{{DOCIR_NS}}}run'):
                            text = run_elem.text or ''
                            run = para.add_run(text)
                            apply_run_style(run, run_elem, region_elem)
                    
                    if len(cell.paragraphs) > 1 and not cell.paragraphs[0].text:
                        p = cell.paragraphs[0]._element
                        p.getparent().remove(p)
                
                # Apply per-cell borders if present
                apply_cell_borders(cell_elem, cell)
                
                col_idx += col_span
            
            row_idx += 1
    
    return top_from_page_top_pt + h_pt


def add_anchored_image(paragraph, image_path: str, left_pt: float, top_pt: float, width_pt: float, height_pt: float):
    """
    Add an image with absolute positioning using wp:anchor.
    
    This creates a floating image anchored to the page with precise coordinates,
    similar to how pdf2docx positions images.
    
    Args:
        paragraph: The paragraph to add the image to
        image_path: Path to the image file
        left_pt: Left position from page left edge (in points)
        top_pt: Top position from page top edge (in points)
        width_pt: Image width (in points)
        height_pt: Image height (in points)
    """
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Emu
    import hashlib
    
    # Convert points to EMU (1 point = 12700 EMU)
    left_emu = int(left_pt * 12700)
    top_emu = int(top_pt * 12700)
    width_emu = int(width_pt * 12700)
    height_emu = int(height_pt * 12700)
    
    # Generate a unique ID for the image
    with open(image_path, 'rb') as f:
        img_hash = hashlib.md5(f.read()).hexdigest()[:8]
    image_id = int(img_hash, 16) % 100000
    
    # Get the run element
    run = paragraph.add_run()
    run_elem = run._r
    
    # First add the picture inline to create the relationship
    inline_shape = run.add_picture(image_path, width=Emu(width_emu), height=Emu(height_emu))
    
    # Find the w:drawing element that was created
    drawing = run_elem.find(qn('w:drawing'))
    if drawing is None:
        return
    
    # Find wp:inline inside w:drawing
    wp_inline = drawing.find(qn('wp:inline'))
    if wp_inline is None:
        return
    
    # Create wp:anchor element
    anchor = OxmlElement('wp:anchor')
    anchor.set('distT', '0')
    anchor.set('distB', '0')
    anchor.set('distL', '0')
    anchor.set('distR', '0')
    anchor.set('simplePos', '0')
    anchor.set('relativeHeight', '0')
    anchor.set('behindDoc', '0')
    anchor.set('locked', '0')
    anchor.set('layoutInCell', '1')
    anchor.set('allowOverlap', '1')
    
    # wp:simplePos (not used but required)
    simplePos = OxmlElement('wp:simplePos')
    simplePos.set('x', '0')
    simplePos.set('y', '0')
    anchor.append(simplePos)
    
    # wp:positionH (horizontal position relative to page)
    positionH = OxmlElement('wp:positionH')
    positionH.set('relativeFrom', 'page')
    posOffset = OxmlElement('wp:posOffset')
    posOffset.text = str(left_emu)
    positionH.append(posOffset)
    anchor.append(positionH)
    
    # wp:positionV (vertical position relative to page)
    positionV = OxmlElement('wp:positionV')
    positionV.set('relativeFrom', 'page')
    posOffset = OxmlElement('wp:posOffset')
    posOffset.text = str(top_emu)
    positionV.append(posOffset)
    anchor.append(positionV)
    
    # wp:extent (image dimensions)
    extent = OxmlElement('wp:extent')
    extent.set('cx', str(width_emu))
    extent.set('cy', str(height_emu))
    anchor.append(extent)
    
    # wp:effectExtent
    effectExtent = OxmlElement('wp:effectExtent')
    effectExtent.set('l', '0')
    effectExtent.set('t', '0')
    effectExtent.set('r', '0')
    effectExtent.set('b', '0')
    anchor.append(effectExtent)
    
    # wp:wrapNone (no text wrapping)
    wrapNone = OxmlElement('wp:wrapNone')
    anchor.append(wrapNone)
    
    # wp:docPr (document properties)
    docPr = OxmlElement('wp:docPr')
    docPr.set('id', str(image_id))
    docPr.set('name', f'Picture {image_id}')
    anchor.append(docPr)
    
    # Copy a:graphic from inline to anchor
    graphic = wp_inline.find(qn('a:graphic'))
    if graphic is not None:
        # Deep copy the graphic element
        graphic_copy = etree.fromstring(etree.tostring(graphic))
        anchor.append(graphic_copy)
    
    # Replace wp:inline with wp:anchor in w:drawing
    drawing.remove(wp_inline)
    drawing.append(anchor)


def add_positioned_image_region(doc: Document, region_elem, page_height_pt: float, current_y_pt: float, assets_dict: dict) -> float:
    """
    Add an image region with absolute positioning using wp:anchor.
    
    Returns the new current_y_pt (bottom of this region).
    """
    bbox = get_bbox(region_elem)
    if bbox is None:
        return current_y_pt
    
    x_pt, y_pt, w_pt, h_pt = bbox
    top_from_page_top_pt = page_height_pt - y_pt
    left_pt = x_pt
    
    # Calculate space needed before this region
    space_before_pt = top_from_page_top_pt - current_y_pt
    if space_before_pt < 0:
        space_before_pt = 0
    
    image_content = region_elem.find(f'{{{DOCIR_NS}}}image_content')
    if image_content is None:
        return current_y_pt
    
    image_ref = image_content.find(f'{{{DOCIR_NS}}}image_reference')
    if image_ref is None:
        return current_y_pt
    
    asset_id = image_ref.get('asset_id')
    if not asset_id or asset_id not in assets_dict:
        # Placeholder
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(space_before_pt)
        para.paragraph_format.left_indent = Pt(left_pt)
        run = para.add_run(f'[Image not found: {asset_id}]')
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        return top_from_page_top_pt + h_pt
    
    asset_info = assets_dict[asset_id]
    image_path = asset_info.get('file_path')
    
    if not image_path or not Path(image_path).exists():
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(space_before_pt)
        para.paragraph_format.left_indent = Pt(left_pt)
        run = para.add_run(f'[Image file not found: {image_path}]')
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        return top_from_page_top_pt + h_pt
    
    # Add image with wp:anchor for precise positioning
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before_pt)
    para.paragraph_format.space_after = Pt(0)
    
    # Use wp:anchor for absolute positioning
    add_anchored_image(para, image_path, left_pt, top_from_page_top_pt, w_pt, h_pt)
    
    # Add caption if present
    caption_elem = image_content.find(f'{{{DOCIR_NS}}}caption')
    if caption_elem is not None and caption_elem.text:
        caption_para = doc.add_paragraph()
        caption_para.paragraph_format.space_before = Pt(2)
        caption_para.paragraph_format.left_indent = Pt(left_pt)
        caption_run = caption_para.add_run(caption_elem.text)
        caption_run.font.size = Pt(10)
        caption_run.font.italic = True
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    return top_from_page_top_pt + h_pt


def process_text_region(doc: Document, region_elem):
    """Process a text region and add it to the document (flow mode)."""
    text_content = region_elem.find(f'{{{DOCIR_NS}}}text_content')
    if text_content is None:
        return
    
    for para_elem in text_content.findall(f'{{{DOCIR_NS}}}paragraph'):
        para = doc.add_paragraph()
        
        # Paragraph alignment (para attr > region computed_style > default)
        alignment = para_elem.get('alignment')
        if not alignment and region_elem is not None:
            computed_style = region_elem.find(f'{{{DOCIR_NS}}}computed_style')
            if computed_style is not None:
                align_elem = computed_style.find(f'{{{DOCIR_NS}}}alignment')
                if align_elem is not None:
                    alignment = align_elem.get('value')
        if alignment:
            set_paragraph_alignment(para, alignment)
        
        for run_elem in para_elem.findall(f'{{{DOCIR_NS}}}run'):
            text = run_elem.text or ''
            run = para.add_run(text)
            apply_run_style(run, run_elem, region_elem)


def process_table_region(doc: Document, region_elem):
    """Process a table region and add it to the document (flow mode)."""
    table_content = region_elem.find(f'{{{DOCIR_NS}}}table_content')
    if table_content is None:
        return
    
    num_rows = int(table_content.get('rows', '0'))
    num_cols = int(table_content.get('cols', '0'))
    
    if num_rows == 0 or num_cols == 0:
        return
    
    table_style_elem = table_content.find(f'{{{DOCIR_NS}}}table_style')
    border_visible = True
    border_color = "000000"
    has_header = False
    
    if table_style_elem is not None:
        border_visible = table_style_elem.get('border_visible', 'true').lower() == 'true'
        border_color_hex = table_style_elem.get('border_color', '#000000')
        border_color = border_color_hex.lstrip('#')
        has_header = table_style_elem.get('header_row', 'false').lower() == 'true'
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    
    if border_visible:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
        borders = OxmlElement('w:tblBorders')
        
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), border_color)
            borders.append(border)
        
        tblPr.append(borders)
    
    # Helper: apply per-cell borders from cell_borders element
    def apply_cell_borders_flow(cell_elem, cell):
        """Apply per-cell border styling from DocIR cell_borders element."""
        cell_borders_elem = cell_elem.find(f'{{{DOCIR_NS}}}cell_borders')
        if cell_borders_elem is None:
            return
        
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        
        tc = cell._tc
        tcPr = tc.tcPr if tc.tcPr is not None else OxmlElement('w:tcPr')
        if tc.tcPr is None:
            tc.insert(0, tcPr)
        
        tcBorders = OxmlElement('w:tcBorders')
        has_borders = False
        
        for edge in ['top', 'bottom', 'left', 'right']:
            border_elem = cell_borders_elem.find(f'{{{DOCIR_NS}}}border_{edge}')
            if border_elem is not None:
                has_borders = True
                border = OxmlElement(f'w:{edge}')
                style = border_elem.get('style', 'single')
                border.set(qn('w:val'), style)
                width_pt = float(border_elem.get('width', '0.5'))
                border.set(qn('w:sz'), str(int(width_pt * 8)))
                border.set(qn('w:space'), '0')
                color = border_elem.get('color', '#000000').lstrip('#')
                border.set(qn('w:color'), color)
                tcBorders.append(border)
        
        if has_borders:
            tcPr.append(tcBorders)
    
    row_idx = 0
    
    header_group = table_content.find(f'{{{DOCIR_NS}}}row_group[@type="header"]')
    if header_group is not None:
        for row_elem in header_group.findall(f'{{{DOCIR_NS}}}row'):
            if row_idx >= num_rows:
                break
            
            col_idx = 0
            for cell_elem in row_elem.findall(f'{{{DOCIR_NS}}}cell'):
                if col_idx >= num_cols:
                    break
                
                cell = table.cell(row_idx, col_idx)
                
                col_span = int(cell_elem.get('col_span', '1'))
                row_span = int(cell_elem.get('row_span', '1'))
                
                if col_span > 1 or row_span > 1:
                    end_row = min(row_idx + row_span - 1, num_rows - 1)
                    end_col = min(col_idx + col_span - 1, num_cols - 1)
                    cell = table.cell(row_idx, col_idx).merge(table.cell(end_row, end_col))
                
                text_content = cell_elem.find(f'{{{DOCIR_NS}}}text_content')
                if text_content is not None:
                    cell.text = ''
                    for para_elem in text_content.findall(f'{{{DOCIR_NS}}}paragraph'):
                        para = cell.add_paragraph()
                        for run_elem in para_elem.findall(f'{{{DOCIR_NS}}}run'):
                            text = run_elem.text or ''
                            run = para.add_run(text)
                            apply_run_style(run, run_elem, region_elem)
                    
                    if len(cell.paragraphs) > 1 and not cell.paragraphs[0].text:
                        p = cell.paragraphs[0]._element
                        p.getparent().remove(p)
                
                if has_header and row_idx == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.bold = True
                
                # Apply per-cell borders if present
                apply_cell_borders_flow(cell_elem, cell)
                
                col_idx += col_span
            
            row_idx += 1
    
    body_group = table_content.find(f'{{{DOCIR_NS}}}row_group[@type="body"]')
    if body_group is not None:
        for row_elem in body_group.findall(f'{{{DOCIR_NS}}}row'):
            if row_idx >= num_rows:
                break
            
            col_idx = 0
            for cell_elem in row_elem.findall(f'{{{DOCIR_NS}}}cell'):
                if col_idx >= num_cols:
                    break
                
                cell = table.cell(row_idx, col_idx)
                
                col_span = int(cell_elem.get('col_span', '1'))
                row_span = int(cell_elem.get('row_span', '1'))
                
                if col_span > 1 or row_span > 1:
                    end_row = min(row_idx + row_span - 1, num_rows - 1)
                    end_col = min(col_idx + col_span - 1, num_cols - 1)
                    cell = table.cell(row_idx, col_idx).merge(table.cell(end_row, end_col))
                
                text_content = cell_elem.find(f'{{{DOCIR_NS}}}text_content')
                if text_content is not None:
                    cell.text = ''
                    for para_elem in text_content.findall(f'{{{DOCIR_NS}}}paragraph'):
                        para = cell.add_paragraph()
                        for run_elem in para_elem.findall(f'{{{DOCIR_NS}}}run'):
                            text = run_elem.text or ''
                            run = para.add_run(text)
                            apply_run_style(run, run_elem, region_elem)
                    
                    if len(cell.paragraphs) > 1 and not cell.paragraphs[0].text:
                        p = cell.paragraphs[0]._element
                        p.getparent().remove(p)
                
                # Apply per-cell borders if present
                apply_cell_borders_flow(cell_elem, cell)
                
                col_idx += col_span
            
            row_idx += 1
    
    doc.add_paragraph()


def process_image_region(doc: Document, region_elem, assets_dict: dict):
    """Process an image region and add it to the document (flow mode)."""
    image_content = region_elem.find(f'{{{DOCIR_NS}}}image_content')
    if image_content is None:
        return
    
    image_ref = image_content.find(f'{{{DOCIR_NS}}}image_reference')
    if image_ref is None:
        return
    
    asset_id = image_ref.get('asset_id')
    if not asset_id or asset_id not in assets_dict:
        para = doc.add_paragraph()
        run = para.add_run(f'[Image asset not found: {asset_id}]')
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        return
    
    asset_info = assets_dict[asset_id]
    image_path = asset_info.get('file_path')
    
    if not image_path or not Path(image_path).exists():
        para = doc.add_paragraph()
        run = para.add_run(f'[Image file not found: {image_path}]')
        run.font.italic = True
        run.font.color.rgb = RGBColor(128, 128, 128)
        return
    
    bbox = get_bbox(region_elem)
    if bbox is not None:
        _, _, w_pt, h_pt = bbox
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(image_path, width=Pt(w_pt), height=Pt(h_pt))
    else:
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(image_path)
    
    caption_elem = image_content.find(f'{{{DOCIR_NS}}}caption')
    if caption_elem is not None and caption_elem.text:
        caption_para = doc.add_paragraph()
        caption_run = caption_para.add_run(caption_elem.text)
        caption_run.font.size = Pt(10)
        caption_run.font.italic = True
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()


def build_assets_dict(root) -> dict:
    """Build a dictionary of assets from the DocIR XML root."""
    assets_dict = {}
    assets_elem = root.find(f'{{{DOCIR_NS}}}assets')
    if assets_elem is not None:
        for asset_elem in assets_elem.findall(f'{{{DOCIR_NS}}}asset'):
            asset_id = asset_elem.get('id')
            if asset_id:
                file_path_elem = asset_elem.find(f'{{{DOCIR_NS}}}file_path')
                file_path = file_path_elem.text if file_path_elem is not None else None
                assets_dict[asset_id] = {
                    'file_path': file_path,
                    'mime_type': asset_elem.get('mime_type', 'image/jpeg'),
                    'width_px': asset_elem.get('width_px'),
                    'height_px': asset_elem.get('height_px'),
                }
    return assets_dict


def get_page_dimensions(root, page_elem) -> Tuple[float, float]:
    """Get page dimensions from page element or default."""
    page_size = page_elem.find(f'{{{DOCIR_NS}}}page_size')
    if page_size is not None:
        w = float(page_size.get('width_pt', 595.30))
        h = float(page_size.get('height_pt', 841.89))
        return (w, h)
    
    # Fallback to default
    metadata = root.find(f'{{{DOCIR_NS}}}metadata')
    if metadata is not None:
        default_size = metadata.find(f'{{{DOCIR_NS}}}default_page_size')
        if default_size is not None:
            w = float(default_size.get('width_pt', 595.30))
            h = float(default_size.get('height_pt', 841.89))
            return (w, h)
    
    return (595.30, 841.89)  # A4 default


def generate_docx(docir_path: Path, output_path: Path, positioned: bool = False, base_dir: Optional[Path] = None):
    """
    Generate DOCX from DocIR XML.
    
    Args:
        docir_path: Path to DocIR XML file
        output_path: Path to output DOCX file
        positioned: If True, use absolute positioning; if False, use flow layout
        base_dir: Base directory for resolving relative image paths
    """
    # Parse DocIR XML
    tree = etree.parse(str(docir_path))
    root = tree.getroot()
    
    # Build assets dictionary
    assets_dict = build_assets_dict(root)
    
    # Resolve relative image paths
    if base_dir is None:
        base_dir = docir_path.parent
    
    for asset_id, asset_info in assets_dict.items():
        file_path = asset_info.get('file_path')
        if file_path and not Path(file_path).is_absolute():
            resolved = base_dir / file_path
            if resolved.exists():
                asset_info['file_path'] = str(resolved)
    
    # Create new document
    doc = Document()
    
    # Get pages
    pages = root.find(f'{{{DOCIR_NS}}}pages')
    if pages is None:
        print("Warning: No pages found in DocIR XML")
        return
    
    page_elements = pages.findall(f'{{{DOCIR_NS}}}page')
    
    if positioned:
        # Position mode: set up first section with correct page size
        first_page_dims = get_page_dimensions(root, page_elements[0])
        setup_positioned_section(doc.sections[0], first_page_dims[0], first_page_dims[1])
        
        # Add title if present (centered at top)
        metadata = root.find(f'{{{DOCIR_NS}}}metadata')
        if metadata is not None:
            title_elem = metadata.find(f'{{{DOCIR_NS}}}title')
            if title_elem is not None and title_elem.text:
                title_para = doc.add_paragraph()
                title_para.paragraph_format.space_before = Pt(20)
                title_para.paragraph_format.space_after = Pt(12)
                title_run = title_para.add_run(title_elem.text)
                title_run.font.size = Pt(16)
                title_run.font.bold = True
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process each page
        for page_idx, page_elem in enumerate(page_elements):
            page_width_pt, page_height_pt = get_page_dimensions(root, page_elem)
            
            # Add new section for subsequent pages
            if page_idx > 0:
                new_section = doc.add_section()
                setup_positioned_section(new_section, page_width_pt, page_height_pt)
            
            # Track current vertical position (from page top)
            current_y_pt = 0
            
            # If we added a title on first page, account for its height
            if page_idx == 0 and metadata is not None:
                title_elem = metadata.find(f'{{{DOCIR_NS}}}title')
                if title_elem is not None and title_elem.text:
                    current_y_pt = 50  # Approximate title height
            
            # Get regions sorted by order
            regions = page_elem.find(f'{{{DOCIR_NS}}}regions')
            if regions is None:
                continue
            
            region_list = []
            for region_elem in regions.findall(f'{{{DOCIR_NS}}}region'):
                order = int(region_elem.get('order', '0'))
                region_list.append((order, region_elem))
            
            region_list.sort(key=lambda x: x[0])
            
            # Process each region with positioning
            for order, region_elem in region_list:
                region_type = region_elem.get('type')
                
                if region_type == 'text':
                    current_y_pt = add_positioned_text_region(doc, region_elem, page_height_pt, current_y_pt)
                elif region_type == 'table':
                    current_y_pt = add_positioned_table_region(doc, region_elem, page_height_pt, current_y_pt)
                elif region_type == 'image':
                    current_y_pt = add_positioned_image_region(doc, region_elem, page_height_pt, current_y_pt, assets_dict)
    else:
        # Flow mode: original behavior
        metadata = root.find(f'{{{DOCIR_NS}}}metadata')
        if metadata is not None:
            title_elem = metadata.find(f'{{{DOCIR_NS}}}title')
            if title_elem is not None and title_elem.text:
                title_para = doc.add_paragraph()
                title_run = title_para.add_run(title_elem.text)
                title_run.font.size = Pt(16)
                title_run.font.bold = True
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
        
        for page_idx, page_elem in enumerate(page_elements):
            if page_idx > 0:
                doc.add_page_break()
                page_para = doc.add_paragraph()
                page_run = page_para.add_run(f'--- Page {page_idx + 1} ---')
                page_run.font.size = Pt(10)
                page_run.font.italic = True
                page_run.font.color.rgb = RGBColor(128, 128, 128)
                page_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph()
            
            regions = page_elem.find(f'{{{DOCIR_NS}}}regions')
            if regions is None:
                continue
            
            region_list = []
            for region_elem in regions.findall(f'{{{DOCIR_NS}}}region'):
                order = int(region_elem.get('order', '0'))
                region_list.append((order, region_elem))
            
            region_list.sort(key=lambda x: x[0])
            
            for order, region_elem in region_list:
                region_type = region_elem.get('type')
                
                if region_type == 'text':
                    process_text_region(doc, region_elem)
                elif region_type == 'table':
                    process_table_region(doc, region_elem)
                elif region_type == 'image':
                    process_image_region(doc, region_elem, assets_dict)
                else:
                    para = doc.add_paragraph()
                    run = para.add_run(f'[{region_type} region]')
                    run.font.italic = True
                    run.font.color.rgb = RGBColor(128, 128, 128)
    
    # Save document
    doc.save(str(output_path))
    mode_str = "positioned" if positioned else "flow"
    print(f"✓ Generated DOCX ({mode_str}): {output_path}")


def main():
    """CLI entry point."""
    import argparse
    
    parser = argparse.ArgumentParser(
        description='Generate DOCX from DocIR XML'
    )
    parser.add_argument(
        'docir_xml',
        type=Path,
        help='Path to DocIR XML file'
    )
    parser.add_argument(
        '-o', '--output',
        type=Path,
        required=True,
        help='Path to output DOCX file'
    )
    parser.add_argument(
        '--positioned',
        action='store_true',
        help='Use absolute positioning (places elements at their PDF coordinates)'
    )
    parser.add_argument(
        '--base-dir',
        type=Path,
        help='Base directory for resolving relative image paths (default: DocIR XML directory)'
    )
    
    args = parser.parse_args()
    
    if not args.docir_xml.exists():
        print(f"Error: DocIR XML file not found: {args.docir_xml}", file=sys.stderr)
        sys.exit(1)
    
    generate_docx(args.docir_xml, args.output, positioned=args.positioned, base_dir=args.base_dir)


if __name__ == '__main__':
    main()
