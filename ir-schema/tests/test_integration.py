#!/usr/bin/env python3
"""
Integration test for the full DocIR pipeline.

Tests:
1. IR Builder: GLM-OCR JSON → DocIR XML
2. Style Extractor: DocIR XML → Styled DocIR XML
3. DOCX Generator: Styled DocIR XML → DOCX
4. Validator: Validate all XML outputs
"""

import sys
import tempfile
from pathlib import Path

# Add parent directory to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from builder.ir_builder import build_docir
from style_extractor.style_extractor import extract_styles_from_docir, VLMConfig
from docx_generator.docx_generator import generate_docx
from validator.semantic_validator import run_all_validations


def test_full_pipeline():
    """Test the full pipeline: GLM-OCR → DocIR → Style → DOCX."""
    # Paths
    repo_root = Path(__file__).parent.parent.parent
    model_json = repo_root / "output-small" / "doc_123c1995a949_small" / "doc_123c1995a949_small_model.json"
    pdf_path = Path("/Users/admin/.hermes/cache/documents/doc_123c1995a949_small.pdf")
    xsd_path = repo_root / "ir-schema" / "docir-v0.1.0.xsd"
    
    # Check inputs exist
    assert model_json.exists(), f"Model JSON not found: {model_json}"
    assert pdf_path.exists(), f"PDF not found: {pdf_path}"
    assert xsd_path.exists(), f"XSD not found: {xsd_path}"
    
    print(f"✓ Input files found")
    print(f"  Model JSON: {model_json}")
    print(f"  PDF: {pdf_path}")
    print(f"  XSD: {xsd_path}")
    
    # Create temp directory for outputs
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)
        
        # Step 1: Build DocIR
        print(f"\n→ Step 1: Building DocIR XML...")
        docir_xml = tmpdir / "test.docir.xml"
        build_docir(
            model_json_path=model_json,
            pdf_path=pdf_path,
            output_path=docir_xml,
            pipeline_info={
                "title": "Integration Test",
                "source": "glm-ocr-pipeline",
                "detection_model": "PP-DocLayoutV3",
                "ocr_engine": "Ollama glm-ocr:latest",
                "style_extractor": "pending"
            }
        )
        
        assert docir_xml.exists(), f"DocIR XML not created: {docir_xml}"
        print(f"✓ DocIR XML created: {docir_xml}")
        
        # Validate Step 1
        print(f"\n→ Validating DocIR XML (Step 1)...")
        result = run_all_validations(docir_xml, xsd_path=xsd_path, strict=False)
        
        if not result.passed:
            print(f"\n✗ FAILED: DocIR XML validation failed")
            for issue in result.issues:
                print(f"  {issue}")
            sys.exit(1)
        
        print(f"✓ DocIR XML validation passed")
        
        # Step 2: Style Extraction
        print(f"\n→ Step 2: Extracting styles...")
        docir_styled = tmpdir / "test-styled.docir.xml"
        
        vlm_config = VLMConfig(
            api_base="http://localhost:11234/v1",
            api_key="change-me-local-key",
            model="qwen3.6-35b-a3b-q7"
        )
        
        extract_styles_from_docir(
            docir_path=docir_xml,
            pdf_path=pdf_path,
            output_path=docir_styled,
            vlm_config=vlm_config,
            dpi=200,
            region_types=["text"]
        )
        
        assert docir_styled.exists(), f"Styled DocIR XML not created: {docir_styled}"
        print(f"✓ Styled DocIR XML created: {docir_styled}")
        
        # Validate Step 2
        print(f"\n→ Validating Styled DocIR XML (Step 2)...")
        result = run_all_validations(docir_styled, xsd_path=xsd_path, strict=False)
        
        if not result.passed:
            print(f"\n✗ FAILED: Styled DocIR XML validation failed")
            for issue in result.issues:
                print(f"  {issue}")
            sys.exit(1)
        
        print(f"✓ Styled DocIR XML validation passed")
        
        # Step 3: DOCX Generation
        print(f"\n→ Step 3: Generating DOCX...")
        docx_path = tmpdir / "test.docx"
        
        generate_docx(
            docir_path=docir_styled,
            output_path=docx_path
        )
        
        assert docx_path.exists(), f"DOCX not created: {docx_path}"
        print(f"✓ DOCX created: {docx_path}")
        
        # Verify DOCX content
        print(f"\n→ Verifying DOCX content...")
        from docx import Document
        
        doc = Document(str(docx_path))
        
        # Check that we have content
        assert len(doc.paragraphs) > 0, "DOCX has no paragraphs"
        print(f"✓ DOCX has {len(doc.paragraphs)} paragraphs")
        
        # Check that styles were applied
        styled_runs = 0
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.name or run.font.size or run.font.bold or run.font.italic:
                    styled_runs += 1
        
        assert styled_runs > 0, "DOCX has no styled runs"
        print(f"✓ DOCX has {styled_runs} styled runs")
        
        # Summary
        print(f"\n{'='*60}")
        print("✓ Full pipeline test PASSED")
        print(f"{'='*60}")
        print(f"  DocIR XML:     {docir_xml}")
        print(f"  Styled XML:    {docir_styled}")
        print(f"  DOCX:          {docx_path}")
        print(f"  Paragraphs:    {len(doc.paragraphs)}")
        print(f"  Styled runs:   {styled_runs}")
        print(f"{'='*60}")


def test_positioned_docx():
    """Test positioned DOCX generation from existing styled XML."""
    repo_root = Path(__file__).parent.parent
    styled_xml = repo_root / "samples" / "small-pdf-styled.xml"
    
    assert styled_xml.exists(), f"Styled XML not found: {styled_xml}"
    
    print(f"\n{'='*60}")
    print("Testing positioned DOCX generation...")
    print(f"{'='*60}")
    
    with tempfile.TemporaryDirectory() as tmpdir:
        tmpdir = Path(tmpdir)
        docx_path = tmpdir / "test-positioned.docx"
        
        # Generate positioned DOCX
        base_dir = repo_root.parent / "output-small" / "doc_123c1995a949_small"
        generate_docx(
            docir_path=styled_xml,
            output_path=docx_path,
            positioned=True,
            base_dir=base_dir
        )
        
        assert docx_path.exists(), f"Positioned DOCX not created: {docx_path}"
        
        # Verify positioned DOCX properties
        from docx import Document
        from docx.shared import Pt
        
        doc = Document(str(docx_path))
        section = doc.sections[0]
        
        # Check page dimensions match PDF (A4: 595.3pt x 841.9pt)
        page_width_pt = section.page_width / 12700
        page_height_pt = section.page_height / 12700
        assert abs(page_width_pt - 595.3) < 1.0, f"Page width mismatch: {page_width_pt:.1f}pt"
        assert abs(page_height_pt - 841.9) < 1.0, f"Page height mismatch: {page_height_pt:.1f}pt"
        print(f"✓ Page dimensions correct: {page_width_pt:.1f}pt × {page_height_pt:.1f}pt")
        
        # Check margins are zero
        assert section.top_margin == 0, f"Top margin not zero: {section.top_margin}"
        assert section.left_margin == 0, f"Left margin not zero: {section.left_margin}"
        print(f"✓ Margins are zero (absolute positioning)")
        
        # Check that paragraphs have positioning
        positioned_paras = 0
        for para in doc.paragraphs:
            pf = para.paragraph_format
            if pf.left_indent is not None or (pf.space_before is not None and pf.space_before > 0):
                positioned_paras += 1
        
        assert positioned_paras > 0, "No positioned paragraphs found"
        print(f"✓ {positioned_paras} paragraphs have positioning")
        
        # Verify specific region positioning
        # Title region: bbox x=217.88pt should give left_indent ≈ 217.88pt
        title_found = False
        for para in doc.paragraphs:
            if "This is Title" in para.text:
                li_pt = para.paragraph_format.left_indent / 12700 if para.paragraph_format.left_indent else 0
                assert abs(li_pt - 217.88) < 2.0, f"Title left_indent mismatch: {li_pt:.1f}pt (expected ~217.88pt)"
                title_found = True
                print(f"✓ Title positioned at x={li_pt:.1f}pt (expected 217.88pt)")
                break
        
        assert title_found, "Title text not found in positioned DOCX"
        
        print(f"\n{'='*60}")
        print("✓ Positioned DOCX test PASSED")
        print(f"{'='*60}")


if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="Run integration tests")
    parser.add_argument("--positioned-only", action="store_true", help="Run only positioned DOCX test")
    parser.add_argument("--full", action="store_true", help="Run full pipeline test (requires VLM)")
    args = parser.parse_args()
    
    if args.positioned_only:
        test_positioned_docx()
    elif args.full:
        test_full_pipeline()
    else:
        # Default: run positioned test (fast, no VLM needed)
        test_positioned_docx()
