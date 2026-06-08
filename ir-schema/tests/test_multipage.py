#!/usr/bin/env python3
"""
Multi-page pipeline test.

Tests that the DocIR pipeline correctly handles multi-page PDFs:
1. IR Builder creates correct page structure
2. DOCX Generator produces correct sections/page breaks
3. Both positioned and flow modes work
"""

import json
import sys
import tempfile
from pathlib import Path

# Add parent to path
sys.path.insert(0, str(Path(__file__).parent.parent))

from builder.ir_builder import build_docir
from docx_generator.docx_generator import generate_docx


def create_multipage_model_json(num_pages: int, regions_per_page: int = 5) -> Path:
    """Create a synthetic multi-page model JSON for testing."""
    pages = []
    for page_idx in range(num_pages):
        regions = []
        for r_idx in range(regions_per_page):
            region = {
                "index": r_idx,
                "label": "text",
                "content": f"Page {page_idx + 1} - Region {r_idx + 1}",
                "bbox_2d": [100, 800 - r_idx * 150, 500, 750 - r_idx * 150],
                "polygon": [
                    [100, 800 - r_idx * 150], [500, 800 - r_idx * 150],
                    [500, 750 - r_idx * 150], [100, 750 - r_idx * 150]
                ]
            }
            regions.append(region)
        pages.append(regions)
    
    output = Path(tempfile.mktemp(suffix='_model.json'))
    with open(output, 'w') as f:
        json.dump(pages, f)
    return output


def create_multipage_pdf(num_pages: int) -> Path:
    """Create a synthetic multi-page PDF for testing."""
    import pymupdf
    
    dst = pymupdf.open()
    # Create blank A4 pages
    for _ in range(num_pages):
        dst.new_page(width=595, height=842)  # A4 in points
    
    output = Path(tempfile.mktemp(suffix='.pdf'))
    dst.save(str(output))
    dst.close()
    return output


def test_ir_builder_multipage():
    """Test IR Builder handles multi-page model JSON correctly."""
    print("Test 1: IR Builder multi-page handling")
    
    num_pages = 5
    regions_per_page = 4
    
    model_json = create_multipage_model_json(num_pages, regions_per_page)
    pdf_path = create_multipage_pdf(num_pages)
    output_xml = Path(tempfile.mktemp(suffix='.docir.xml'))
    
    try:
        pipeline_info = {
            "title": "Multi-Page Test",
            "source": "test",
            "detection_model": "test",
            "ocr_engine": "test",
            "style_extractor": "none"
        }
        
        build_docir(model_json, pdf_path, output_xml, pipeline_info)
        
        # Verify output
        from lxml import etree
        tree = etree.parse(str(output_xml))
        root = tree.getroot()
        
        DOCIR_NS = "urn:docir:v0.1"
        pages = root.find(f'{{{DOCIR_NS}}}pages')
        page_elements = pages.findall(f'{{{DOCIR_NS}}}page')
        
        assert len(page_elements) == num_pages, f"Expected {num_pages} pages, got {len(page_elements)}"
        
        for i, page_elem in enumerate(page_elements):
            assert page_elem.get('index') == str(i), f"Page index mismatch: expected {i}, got {page_elem.get('index')}"
            
            regions = page_elem.find(f'{{{DOCIR_NS}}}regions')
            region_list = regions.findall(f'{{{DOCIR_NS}}}region')
            assert len(region_list) == regions_per_page, \
                f"Page {i}: expected {regions_per_page} regions, got {len(region_list)}"
            
            # Verify region IDs include page index
            for region in region_list:
                region_id = region.get('id')
                assert f'_p{i}' in region_id, f"Region ID {region_id} should contain _p{i}"
        
        # Verify metadata
        metadata = root.find(f'{{{DOCIR_NS}}}metadata')
        page_count = metadata.find(f'{{{DOCIR_NS}}}page_count')
        assert page_count.text == str(num_pages), f"page_count mismatch: {page_count.text} != {num_pages}"
        
        print(f"  ✓ {num_pages} pages, {regions_per_page} regions each")
        print(f"  ✓ Region IDs correctly include page index")
        print(f"  ✓ Metadata page_count correct")
        
    finally:
        model_json.unlink(missing_ok=True)
        pdf_path.unlink(missing_ok=True)
        output_xml.unlink(missing_ok=True)


def test_docx_generator_multipage_positioned():
    """Test DOCX Generator produces correct sections for multi-page in positioned mode."""
    print("\nTest 2: DOCX Generator multi-page positioned mode")
    
    num_pages = 4
    regions_per_page = 3
    
    model_json = create_multipage_model_json(num_pages, regions_per_page)
    pdf_path = create_multipage_pdf(num_pages)
    docir_xml = Path(tempfile.mktemp(suffix='.docir.xml'))
    docx_output = Path(tempfile.mktemp(suffix='.docx'))
    
    try:
        # Build DocIR
        pipeline_info = {
            "title": "Multi-Page Positioned Test",
            "source": "test",
            "detection_model": "test",
            "ocr_engine": "test",
            "style_extractor": "none"
        }
        build_docir(model_json, pdf_path, docir_xml, pipeline_info)
        
        # Generate DOCX
        generate_docx(docir_xml, docx_output, positioned=True)
        
        # Verify
        from docx import Document
        doc = Document(str(docx_output))
        
        assert len(doc.sections) == num_pages, \
            f"Expected {num_pages} sections, got {len(doc.sections)}"
        
        # Verify each section has correct page size (A4)
        for i, section in enumerate(doc.sections):
            width_in = section.page_width.inches
            height_in = section.page_height.inches
            assert abs(width_in - 8.27) < 0.1, f"Section {i} width unexpected: {width_in}"
            assert abs(height_in - 11.69) < 0.1, f"Section {i} height unexpected: {height_in}"
        
        print(f"  ✓ {num_pages} sections created")
        print(f"  ✓ All sections have correct A4 page size")
        
    finally:
        model_json.unlink(missing_ok=True)
        pdf_path.unlink(missing_ok=True)
        docir_xml.unlink(missing_ok=True)
        docx_output.unlink(missing_ok=True)


def test_docx_generator_multipage_flow():
    """Test DOCX Generator produces page breaks in flow mode."""
    print("\nTest 3: DOCX Generator multi-page flow mode")
    
    num_pages = 3
    regions_per_page = 2
    
    model_json = create_multipage_model_json(num_pages, regions_per_page)
    pdf_path = create_multipage_pdf(num_pages)
    docir_xml = Path(tempfile.mktemp(suffix='.docir.xml'))
    docx_output = Path(tempfile.mktemp(suffix='.docx'))
    
    try:
        pipeline_info = {
            "title": "Multi-Page Flow Test",
            "source": "test",
            "detection_model": "test",
            "ocr_engine": "test",
            "style_extractor": "none"
        }
        build_docir(model_json, pdf_path, docir_xml, pipeline_info)
        
        # Generate DOCX in flow mode (no --positioned)
        generate_docx(docir_xml, docx_output, positioned=False)
        
        # Verify page break markers
        from docx import Document
        doc = Document(str(docx_output))
        
        page_markers = [p.text for p in doc.paragraphs if p.text and '--- Page' in p.text]
        expected_markers = [f'--- Page {i+1} ---' for i in range(1, num_pages)]
        
        assert len(page_markers) == num_pages - 1, \
            f"Expected {num_pages - 1} page markers, got {len(page_markers)}"
        
        for marker in expected_markers:
            assert marker in page_markers, f"Missing page marker: {marker}"
        
        print(f"  ✓ {num_pages - 1} page break markers found")
        print(f"  ✓ Markers: {page_markers}")
        
    finally:
        model_json.unlink(missing_ok=True)
        pdf_path.unlink(missing_ok=True)
        docir_xml.unlink(missing_ok=True)
        docx_output.unlink(missing_ok=True)


if __name__ == '__main__':
    print("=" * 60)
    print("Multi-Page Pipeline Tests")
    print("=" * 60)
    
    test_ir_builder_multipage()
    test_docx_generator_multipage_positioned()
    test_docx_generator_multipage_flow()
    
    print("\n" + "=" * 60)
    print("✓ All multi-page tests passed!")
    print("=" * 60)
