#!/usr/bin/env python3
"""
Test DOCX translator extraction and application logic.
"""

import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from docx.shared import Pt, RGBColor
from translator.translator import extract_text_segments, apply_translations, TextSegment


def test_extraction_and_application():
    """Test that we can extract text segments and apply translations while preserving formatting."""
    
    # Create a test DOCX with various formatting
    doc = Document()
    
    # Paragraph with formatted text
    para = doc.add_paragraph()
    run1 = para.add_run("Hello ")
    run1.font.bold = True
    run1.font.size = Pt(14)
    
    run2 = para.add_run("World")
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    
    # Table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Cell 1"
    table.cell(1, 1).text = "Cell 2"
    
    # Save and reload
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        input_path = Path(f.name)
    doc.save(str(input_path))
    
    # Reload
    doc = Document(str(input_path))
    
    # Extract segments
    segments = extract_text_segments(doc)
    print(f"Extracted {len(segments)} segments:")
    for seg in segments:
        location = "table" if seg.is_in_table else "body"
        print(f"  [{location}] p{seg.paragraph_idx}:r{seg.run_idx} = '{seg.text}'")
    
    # Verify extraction
    assert len(segments) == 6, f"Expected 6 segments, got {len(segments)}"
    
    body_segments = [s for s in segments if not s.is_in_table]
    table_segments = [s for s in segments if s.is_in_table]
    
    assert len(body_segments) == 2, f"Expected 2 body segments, got {len(body_segments)}"
    assert len(table_segments) == 4, f"Expected 4 table segments, got {len(table_segments)}"
    
    # Verify body text
    assert body_segments[0].text == "Hello "
    assert body_segments[1].text == "World"
    
    # Apply fake translations
    for seg in segments:
        if seg.text == "Hello ":
            seg.translated = "你好 "
        elif seg.text == "World":
            seg.translated = "世界"
        elif seg.text == "Header 1":
            seg.translated = "标题 1"
        elif seg.text == "Header 2":
            seg.translated = "标题 2"
        elif seg.text == "Cell 1":
            seg.translated = "单元格 1"
        elif seg.text == "Cell 2":
            seg.translated = "单元格 2"
    
    # Apply translations
    doc = apply_translations(doc, segments)
    
    # Save and reload
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        output_path = Path(f.name)
    doc.save(str(output_path))
    
    # Verify formatting preserved
    doc = Document(str(output_path))
    
    # Check body paragraph
    para = doc.paragraphs[0]
    assert len(para.runs) == 2, f"Expected 2 runs, got {len(para.runs)}"
    
    run1 = para.runs[0]
    assert run1.text == "你好 ", f"Expected '你好 ', got '{run1.text}'"
    assert run1.font.bold == True, "Bold not preserved"
    assert run1.font.size == Pt(14), "Font size not preserved"
    
    run2 = para.runs[1]
    assert run2.text == "世界", f"Expected '世界', got '{run2.text}'"
    assert run2.font.italic == True, "Italic not preserved"
    assert run2.font.color.rgb == RGBColor(0xFF, 0x00, 0x00), "Color not preserved"
    
    # Check table
    table = doc.tables[0]
    assert table.cell(0, 0).text == "标题 1"
    assert table.cell(0, 1).text == "标题 2"
    assert table.cell(1, 0).text == "单元格 1"
    assert table.cell(1, 1).text == "单元格 2"
    
    # Cleanup
    input_path.unlink()
    output_path.unlink()
    
    print()
    print("=" * 60)
    print("✓ ALL TRANSLATOR TESTS PASSED")
    print("=" * 60)
    print()
    print("Verified:")
    print("  ✓ Text extraction from paragraphs and tables")
    print("  ✓ Translation application preserves formatting")
    print("  ✓ Bold, italic, font size, color all preserved")
    print("  ✓ Table cell translations work correctly")


if __name__ == '__main__':
    test_extraction_and_application()
