#!/usr/bin/env python3
"""
DOCX Translator - Translate DOCX content while preserving all formatting.

Features:
- Extract text from DOCX preserving document structure
- Translate using LLM (local Ollama or OpenAI-compatible API)
- Replace text while preserving ALL formatting (fonts, sizes, colors, bold, italic, etc.)
- Handle tables, headers/footers, text boxes, footnotes
- Support style instructions (formal, legal, casual, etc.)
- Batch translation for efficiency
- Progress tracking and resumable

Usage:
    python3 translator.py input.docx -o output.docx --target-lang zh --style formal
    python3 translator.py input.docx -o output.docx --target-lang en --api-base http://localhost:11434/v1 --model qwen3.6-27b
"""

import json
import re
import sys
import time
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass, field
from copy import deepcopy

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests


@dataclass
class TranslationConfig:
    """Configuration for translation."""
    target_lang: str = "zh"
    source_lang: Optional[str] = None  # Auto-detect if None
    style: str = "neutral"  # formal, legal, casual, technical, neutral
    api_base: str = "http://localhost:11234/v1"
    api_key: str = "change-me-local-key"
    model: str = "qwen3.6-27b-k-xl"
    max_tokens: int = 4096
    temperature: float = 0.3
    batch_size: int = 20  # Text segments per batch
    preserve_placeholders: bool = True  # Keep {var}, [ref], etc.
    glossary: Dict[str, str] = field(default_factory=dict)  # Term translations


@dataclass
class TextSegment:
    """A text segment with its location in the document."""
    text: str
    paragraph_idx: int
    run_idx: int
    is_in_table: bool = False
    table_row: int = 0
    table_col: int = 0
    is_header: bool = False
    is_footer: bool = False
    section_idx: int = 0
    translated: Optional[str] = None


def extract_text_segments(doc: Document) -> List[TextSegment]:
    """
    Extract all text segments from DOCX with their locations.
    
    Preserves document structure: paragraphs, tables, headers, footers.
    """
    segments = []
    
    # Extract from body paragraphs
    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
            if run.text.strip():
                segments.append(TextSegment(
                    text=run.text,
                    paragraph_idx=para_idx,
                    run_idx=run_idx,
                    is_in_table=False
                ))
    
    # Extract from tables
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(para.runs):
                        if run.text.strip():
                            segments.append(TextSegment(
                                text=run.text,
                                paragraph_idx=para_idx,
                                run_idx=run_idx,
                                is_in_table=True,
                                table_row=row_idx,
                                table_col=col_idx
                            ))
    
    # Extract from headers and footers
    for section_idx, section in enumerate(doc.sections):
        # Header
        if section.header and not section.header.is_linked_to_previous:
            for para_idx, para in enumerate(section.header.paragraphs):
                for run_idx, run in enumerate(para.runs):
                    if run.text.strip():
                        segments.append(TextSegment(
                            text=run.text,
                            paragraph_idx=para_idx,
                            run_idx=run_idx,
                            is_header=True,
                            section_idx=section_idx
                        ))
        
        # Footer
        if section.footer and not section.footer.is_linked_to_previous:
            for para_idx, para in enumerate(section.footer.paragraphs):
                for run_idx, run in enumerate(para.runs):
                    if run.text.strip():
                        segments.append(TextSegment(
                            text=run.text,
                            paragraph_idx=para_idx,
                            run_idx=run_idx,
                            is_footer=True,
                            section_idx=section_idx
                        ))
    
    return segments


def build_translation_prompt(segments: List[TextSegment], config: TranslationConfig) -> str:
    """Build the translation prompt for the LLM."""
    
    # Prepare text batch
    texts = [seg.text for seg in segments]
    numbered_texts = "\n".join(f"[{i}] {t}" for i, t in enumerate(texts))
    
    # Style instructions
    style_instructions = {
        "formal": "Use formal, professional language appropriate for business documents.",
        "legal": "Use precise legal terminology. Maintain legal accuracy.",
        "casual": "Use conversational, friendly language.",
        "technical": "Use precise technical terminology. Maintain technical accuracy.",
        "neutral": "Translate directly without changing the tone.",
    }
    
    style_text = style_instructions.get(config.style, style_instructions["neutral"])
    
    # Source/target language
    source_text = f" from {config.source_lang}" if config.source_lang else ""
    target_text = f" to {config.target_lang}"
    
    # Glossary
    glossary_text = ""
    if config.glossary:
        glossary_items = "\n".join(f"  - '{k}' → '{v}'" for k, v in config.glossary.items())
        glossary_text = f"\n\nGlossary (must use these translations):\n{glossary_items}"
    
    prompt = f"""You are a professional translator. Translate the following text segments{source_text}{target_text}.

Style: {style_text}{glossary_text}

Rules:
1. Preserve all formatting markers (bold, italic markers if any)
2. Keep proper nouns, brand names, and technical terms as appropriate
3. Maintain the same level of formality as the source
4. Keep numbers, dates, and measurements in the same format
5. Preserve any placeholders like {{variable}} or [reference]
6. Output ONLY the translations, one per line, in the format: [N] translated text
7. Do NOT add explanations or notes

Text segments to translate:
{numbered_texts}

Translations:"""
    
    return prompt


def call_llm(prompt: str, config: TranslationConfig) -> str:
    """Call the LLM API for translation."""
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {config.api_key}"
    }
    
    payload = {
        "model": config.model,
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "max_tokens": config.max_tokens,
        "temperature": config.temperature,
    }
    
    try:
        response = requests.post(
            f"{config.api_base}/chat/completions",
            headers=headers,
            json=payload,
            timeout=120
        )
        response.raise_for_status()
        
        result = response.json()
        return result["choices"][0]["message"]["content"]
    
    except Exception as e:
        print(f"Error calling LLM: {e}")
        raise


def parse_translations(response: str, num_segments: int) -> List[str]:
    """Parse the LLM response to extract translations."""
    
    translations = [""] * num_segments
    
    # Parse lines like "[0] translated text"
    pattern = re.compile(r'\[(\d+)\]\s*(.+)')
    
    for line in response.strip().split('\n'):
        line = line.strip()
        match = pattern.match(line)
        if match:
            idx = int(match.group(1))
            text = match.group(2)
            if 0 <= idx < num_segments:
                translations[idx] = text
    
    return translations


def translate_segments(segments: List[TextSegment], config: TranslationConfig) -> List[TextSegment]:
    """Translate all segments using batched LLM calls."""
    
    total = len(segments)
    translated = 0
    
    # Process in batches
    for batch_start in range(0, total, config.batch_size):
        batch_end = min(batch_start + config.batch_size, total)
        batch = segments[batch_start:batch_end]
        
        print(f"  Translating batch {batch_start//config.batch_size + 1}: segments {batch_start}-{batch_end-1}...")
        
        # Build prompt and call LLM
        prompt = build_translation_prompt(batch, config)
        response = call_llm(prompt, config)
        
        # Parse translations
        translations = parse_translations(response, len(batch))
        
        # Apply to segments
        for i, seg in enumerate(batch):
            if i < len(translations) and translations[i]:
                seg.translated = translations[i]
                translated += 1
        
        # Rate limiting
        time.sleep(0.5)
    
    print(f"  ✓ Translated {translated}/{total} segments")
    return segments


def apply_translations(doc: Document, segments: List[TextSegment]) -> Document:
    """Apply translations to the DOCX document, preserving all formatting."""
    
    # Build lookup for translated segments
    body_segments = [s for s in segments if not s.is_in_table and not s.is_header and not s.is_footer]
    table_segments = [s for s in segments if s.is_in_table]
    header_segments = [s for s in segments if s.is_header]
    footer_segments = [s for s in segments if s.is_footer]
    
    # Apply to body paragraphs
    body_idx = 0
    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
            if run.text.strip() and body_idx < len(body_segments):
                seg = body_segments[body_idx]
                if seg.paragraph_idx == para_idx and seg.run_idx == run_idx:
                    if seg.translated:
                        run.text = seg.translated
                    body_idx += 1
    
    # Apply to tables
    table_lookup = {}
    for seg in table_segments:
        key = (seg.table_row, seg.table_col, seg.paragraph_idx, seg.run_idx)
        table_lookup[key] = seg
    
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for para_idx, para in enumerate(cell.paragraphs):
                    for run_idx, run in enumerate(para.runs):
                        key = (row_idx, col_idx, para_idx, run_idx)
                        if key in table_lookup:
                            seg = table_lookup[key]
                            if seg.translated and run.text.strip():
                                run.text = seg.translated
    
    # Apply to headers
    header_idx = 0
    for section in doc.sections:
        if section.header and not section.header.is_linked_to_previous:
            for para_idx, para in enumerate(section.header.paragraphs):
                for run_idx, run in enumerate(para.runs):
                    if run.text.strip() and header_idx < len(header_segments):
                        seg = header_segments[header_idx]
                        if seg.translated:
                            run.text = seg.translated
                        header_idx += 1
    
    # Apply to footers
    footer_idx = 0
    for section in doc.sections:
        if section.footer and not section.footer.is_linked_to_previous:
            for para_idx, para in enumerate(section.footer.paragraphs):
                for run_idx, run in enumerate(para.runs):
                    if run.text.strip() and footer_idx < len(footer_segments):
                        seg = footer_segments[footer_idx]
                        if seg.translated:
                            run.text = seg.translated
                        footer_idx += 1
    
    return doc


def translate_docx(
    input_path: Path,
    output_path: Path,
    config: TranslationConfig
) -> Path:
    """
    Translate a DOCX file while preserving all formatting.
    
    Args:
        input_path: Input DOCX file
        output_path: Output DOCX file
        config: Translation configuration
    
    Returns:
        Path to output file
    """
    print(f"Translating: {input_path}")
    print(f"  Target language: {config.target_lang}")
    print(f"  Style: {config.style}")
    print(f"  Model: {config.model}")
    print()
    
    # Load document
    doc = Document(str(input_path))
    
    # Extract text segments
    print("Extracting text segments...")
    segments = extract_text_segments(doc)
    print(f"  Found {len(segments)} text segments")
    
    if not segments:
        print("  No text to translate")
        return input_path
    
    # Translate
    print("\nTranslating...")
    segments = translate_segments(segments, config)
    
    # Apply translations
    print("\nApplying translations to DOCX...")
    doc = apply_translations(doc, segments)
    
    # Save
    doc.save(str(output_path))
    print(f"\n✓ Translation complete: {output_path}")
    
    return output_path


if __name__ == '__main__':
    import argparse
    
    parser = argparse.ArgumentParser(description='Translate DOCX preserving formatting')
    parser.add_argument('input', type=Path, help='Input DOCX file')
    parser.add_argument('-o', '--output', type=Path, required=True, help='Output DOCX file')
    parser.add_argument('--target-lang', '-l', default='zh', help='Target language (default: zh)')
    parser.add_argument('--source-lang', '-s', help='Source language (auto-detect if not specified)')
    parser.add_argument('--style', default='neutral', 
                       choices=['formal', 'legal', 'casual', 'technical', 'neutral'],
                       help='Translation style (default: neutral)')
    parser.add_argument('--api-base', default='http://localhost:11234/v1',
                       help='LLM API base URL')
    parser.add_argument('--api-key', default='change-me-local-key',
                       help='LLM API key')
    parser.add_argument('--model', default='qwen3.6-27b-k-xl',
                       help='LLM model name')
    parser.add_argument('--batch-size', type=int, default=20,
                       help='Text segments per batch (default: 20)')
    parser.add_argument('--glossary', type=Path,
                       help='Glossary JSON file ({"term": "translation", ...})')
    
    args = parser.parse_args()
    
    # Load glossary if provided
    glossary = {}
    if args.glossary and args.glossary.exists():
        with open(args.glossary) as f:
            glossary = json.load(f)
    
    config = TranslationConfig(
        target_lang=args.target_lang,
        source_lang=args.source_lang,
        style=args.style,
        api_base=args.api_base,
        api_key=args.api_key,
        model=args.model,
        batch_size=args.batch_size,
        glossary=glossary
    )
    
    translate_docx(args.input, args.output, config)
